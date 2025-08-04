#!/usr/bin/env python3
"""
🚌 Voigts Student Data Processing Suite - Complete Transportation Data Management
A comprehensive web interface for managing the entire student transportation data workflow.

Features:
- Word to PDF Conversion for transportation forms
- PDF Data Extraction with AI for student preferences
- Data Validation and Comparison with district databases  
- Traversa-Ready File Generation with transportation analysis
- All-in-One Automated Processing
- Modern, responsive UI with Voigts branding
"""

import streamlit as st
import pandas as pd
import io
import os
import tempfile
from datetime import datetime
import base64
from pathlib import Path
import zipfile
import logging

# Import student data comparator if available
try:
    from student_data_comparator import StudentDataComparator
    COMPARATOR_AVAILABLE = True
except ImportError:
    COMPARATOR_AVAILABLE = False

# Import Word to PDF functionality
try:
    from docx import Document
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from io import BytesIO
    DOCX_TO_PDF_AVAILABLE = True
except ImportError:
    DOCX_TO_PDF_AVAILABLE = False
    
# Keep docx2pdf as fallback for older systems
try:
    from docx2pdf import convert as docx_convert
    DOCX2PDF_FALLBACK = True
except ImportError:
    DOCX2PDF_FALLBACK = False
    
# Import AI extraction functionality  
try:
    from azure.ai.formrecognizer import DocumentAnalysisClient
    from azure.core.credentials import AzureKeyCredential
    AI_EXTRACTOR_AVAILABLE = True
except ImportError:
    AI_EXTRACTOR_AVAILABLE = False

# Import Traversa data processor (simplified for cloud deployment)
try:
    from traversa_data_processor import TraversaDataProcessor
    TRAVERSA_PROCESSOR_AVAILABLE = True
except ImportError:
    TRAVERSA_PROCESSOR_AVAILABLE = False

# Configure Streamlit page
st.set_page_config(
    page_title="Voigts Student Data Processing Suite",
    page_icon="🚌",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for modern styling
st.markdown("""
<style>
    /* Glass effect background */
    .stApp {
        background: linear-gradient(135deg, #1e3a8a 0%, #3b82f6 25%, #60a5fa 50%, #fbbf24 75%, #f59e0b 100%);
        background-attachment: fixed;
    }
    
    .main-header {
        font-size: 3rem;
        font-weight: 700;
        text-align: center;
        background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 2rem;
        text-shadow: 0 4px 8px rgba(0,0,0,0.1);
    }
    
    .emoji-icon {
        -webkit-text-fill-color: #1e3a8a !important;
        background: none !important;
        color: #1e3a8a !important;
        margin-right: 0.5rem;
    }
    
    .subtitle {
        text-align: center;
        font-size: 1.2rem;
        color: #1e3c72;
        margin-bottom: 3rem;
        font-weight: 500;
    }
    
    .logo-section {
        background: rgba(255, 255, 255, 0.15);
        backdrop-filter: blur(10px);
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1rem 0 2rem 0;
        border: 1px solid rgba(255, 255, 255, 0.2);
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.1);
    }
    
    .logo-section h2 {
        color: #1e3c72;
        font-size: 1.8rem;
        font-weight: 700;
        margin: 0 0 0.5rem 0;
        text-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
    }
    
    .logo-section p {
        color: #1e3c72;
        font-size: 1rem;
        font-weight: 500;
        margin: 0;
        opacity: 0.9;
    }
        font-weight: 500;
    }
    
    .upload-section {
        background: rgba(255, 255, 255, 0.2);
        backdrop-filter: blur(15px);
        border-radius: 20px;
        padding: 1.5rem;
        margin: 1rem 0;
        border: 1px solid rgba(255, 255, 255, 0.3);
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.1);
        transition: all 0.3s ease;
        min-height: 250px;
        display: flex;
        flex-direction: column;
        justify-content: space-between;
    }
    
    .upload-section:hover {
        background: rgba(255, 255, 255, 0.3);
        border: 1px solid rgba(251, 191, 36, 0.5);
        transform: translateY(-3px);
        box-shadow: 0 12px 40px rgba(0, 0, 0, 0.2);
    }
    
    .results-container {
        background: rgba(255, 255, 255, 0.1);
        backdrop-filter: blur(15px);
        border-radius: 20px;
        padding: 2rem;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.1);
        border: 1px solid rgba(255, 255, 255, 0.2);
        margin: 2rem 0;
    }
    
    .metric-card {
        background: linear-gradient(135deg, rgba(30, 58, 138, 0.9) 0%, rgba(251, 191, 36, 0.9) 100%);
        backdrop-filter: blur(10px);
        color: white;
        padding: 1.5rem;
        border-radius: 15px;
        text-align: center;
        margin: 0.5rem;
        border: 1px solid rgba(255, 255, 255, 0.2);
        box-shadow: 0 4px 16px rgba(0, 0, 0, 0.1);
        transition: all 0.3s ease;
    }
    
    .metric-card:hover {
        transform: translateY(-3px);
        box-shadow: 0 8px 25px rgba(0, 0, 0, 0.2);
    }
    
    .success-box {
        background: rgba(16, 185, 129, 0.15);
        backdrop-filter: blur(10px);
        border: 1px solid rgba(16, 185, 129, 0.3);
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1rem 0;
        color: #064e3b;
        box-shadow: 0 4px 16px rgba(16, 185, 129, 0.1);
    }
    
    .warning-box {
        background: rgba(251, 191, 36, 0.15);
        backdrop-filter: blur(10px);
        border: 1px solid rgba(251, 191, 36, 0.3);
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1rem 0;
        color: #92400e;
        box-shadow: 0 4px 16px rgba(251, 191, 36, 0.1);
    }
    
    .download-button {
        background: linear-gradient(135deg, #10b981 0%, #34d399 100%);
        backdrop-filter: blur(10px);
        color: white;
        border: none;
        padding: 12px 24px;
        border-radius: 12px;
        font-weight: 600;
        cursor: pointer;
        transition: all 0.3s ease;
        box-shadow: 0 4px 16px rgba(16, 185, 129, 0.3);
    }
    
    .download-button:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 25px rgba(16, 185, 129, 0.4);
    }
    
    .stButton > button {
        width: 100% !important;
        background: rgba(255, 255, 255, 0.15) !important;
        backdrop-filter: blur(20px) !important;
        -webkit-backdrop-filter: blur(20px) !important;
        color: #1e3c72 !important;
        border: 1px solid rgba(255, 255, 255, 0.4) !important;
        border-radius: 12px !important;
        padding: 1.5rem 2rem !important;
        font-weight: 600 !important;
        font-size: 1.1rem !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.15) !important;
        min-height: 80px !important;
        max-height: 80px !important;
        white-space: nowrap !important;
        overflow: hidden !important;
        text-overflow: ellipsis !important;
        line-height: 1.4 !important;
        text-align: center !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        margin: 0 0 0.5rem 0 !important;
        box-sizing: border-box !important;
    }
    
    .stButton > button:hover {
        transform: translateY(-3px) !important;
        box-shadow: 0 12px 40px rgba(0, 0, 0, 0.25) !important;
        background: rgba(255, 255, 255, 0.25) !important;
        border: 1px solid rgba(251, 191, 36, 0.6) !important;
        color: #1e3c72 !important;
    }
    
    /* Ensure column alignment */
    [data-testid="stHorizontalBlock"] {
        display: flex !important;
        width: 100% !important;
        gap: 1rem !important;
        margin-bottom: 0 !important;
        clear: both !important;
    }
    
    [data-testid="stHorizontalBlock"] > div {
        display: flex !important;
        align-items: stretch !important;
        height: auto !important;
        width: 100% !important;
    }
    
    [data-testid="stColumn"] {
        display: flex !important;
        flex-direction: column !important;
        height: 100% !important;
        width: 100% !important;
        flex: 1 !important;
    }
    
    [data-testid="stColumn"] > div {
        height: 100% !important;
        display: flex !important;
        flex-direction: column !important;
        width: 100% !important;
    }
    
    /* Ensure buttons fill their containers properly */
    [data-testid="stColumn"] [data-testid="stButton"] {
        width: 100% !important;
        height: 100% !important;
    }
    
    /* Special styling for dashboard navigation buttons */
    [data-testid="stButton"] button[data-baseweb="button"] {
        width: 100% !important;
        background: rgba(255, 255, 255, 0.15) !important;
        backdrop-filter: blur(20px) !important;
        -webkit-backdrop-filter: blur(20px) !important;
        border: 1px solid rgba(255, 255, 255, 0.4) !important;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.15) !important;
        transition: all 0.3s ease !important;
        min-height: 80px !important;
        max-height: 80px !important;
        font-size: 1.1rem !important;
        font-weight: 600 !important;
        text-align: center !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        border-radius: 12px !important;
        color: #1e3c72 !important;
        margin: 0 0 0.5rem 0 !important;
        padding: 1.5rem 2rem !important;
        box-sizing: border-box !important;
        white-space: nowrap !important;
        overflow: hidden !important;
        text-overflow: ellipsis !important;
    }
    
    [data-testid="stButton"] button[data-baseweb="button"]:hover {
        background: rgba(255, 255, 255, 0.25) !important;
        transform: translateY(-3px) !important;
        box-shadow: 0 12px 40px rgba(0, 0, 0, 0.25) !important;
        border: 1px solid rgba(251, 191, 36, 0.6) !important;
        color: #1e3c72 !important;
    }
    
    /* Special styling for the All-in-One button to make it stand out */
    [data-testid="stButton"] button[data-baseweb="button"][aria-label*="All-in-One"] {
        background: linear-gradient(135deg, rgba(251, 191, 36, 0.2) 0%, rgba(30, 60, 114, 0.2) 100%) !important;
        border: 1px solid rgba(251, 191, 36, 0.5) !important;
        font-size: 1.2rem !important;
        font-weight: 700 !important;
        min-height: 90px !important;
        max-height: 90px !important;
    }
    
    [data-testid="stButton"] button[data-baseweb="button"][aria-label*="All-in-One"]:hover {
        background: linear-gradient(135deg, rgba(251, 191, 36, 0.3) 0%, rgba(30, 60, 114, 0.3) 100%) !important;
        border: 1px solid rgba(251, 191, 36, 0.7) !important;
        transform: translateY(-4px) !important;
        box-shadow: 0 16px 50px rgba(251, 191, 36, 0.3) !important;
    }
    
    /* Glass effect for file uploader */
    .stFileUploader {
        background: rgba(255, 255, 255, 0.2) !important;
        backdrop-filter: blur(10px) !important;
        border-radius: 15px !important;
        border: 1px solid rgba(255, 255, 255, 0.3) !important;
        padding: 1rem !important;
        box-shadow: 0 4px 16px rgba(0, 0, 0, 0.1) !important;
    }
    
    /* File uploader dropzone styling */
    .stFileUploader > div > div {
        background: rgba(255, 255, 255, 0.15) !important;
        backdrop-filter: blur(8px) !important;
        border-radius: 12px !important;
        border: 2px dashed rgba(251, 191, 36, 0.5) !important;
        transition: all 0.3s ease !important;
    }
    
    .stFileUploader > div > div:hover {
        background: rgba(255, 255, 255, 0.25) !important;
        border: 2px dashed rgba(251, 191, 36, 0.8) !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 25px rgba(0, 0, 0, 0.15) !important;
    }
    
    /* File uploader button styling */
    .stFileUploader button {
        background: linear-gradient(135deg, #1e3a8a 0%, #fbbf24 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 10px !important;
        padding: 0.75rem 1.5rem !important;
        font-weight: 600 !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 4px 12px rgba(30, 58, 138, 0.3) !important;
    }
    
    .stFileUploader button:hover {
        background: linear-gradient(135deg, #1e40af 0%, #f59e0b 100%) !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px rgba(30, 58, 138, 0.4) !important;
    }
    
    /* Glass effect for sliders */
    .stSlider {
        background: rgba(255, 255, 255, 0.1);
        backdrop-filter: blur(5px);
        border-radius: 10px;
        padding: 1rem;
    }
    
    /* Glass effect for info boxes */
    .stInfo {
        background: rgba(59, 130, 246, 0.1) !important;
        backdrop-filter: blur(10px);
        border: 1px solid rgba(59, 130, 246, 0.2);
        border-radius: 10px;
    }
    
    .stWarning {
        background: rgba(251, 191, 36, 0.1) !important;
        backdrop-filter: blur(10px);
        border: 1px solid rgba(251, 191, 36, 0.2);
        border-radius: 10px;
    }
    
    .stSuccess {
        background: rgba(16, 185, 129, 0.1) !important;
        backdrop-filter: blur(10px);
        border: 1px solid rgba(16, 185, 129, 0.2);
        border-radius: 10px;
    }
    
    /* Hide any empty upload-section divs */
    .upload-section:empty {
        display: none !important;
        height: 0 !important;
        padding: 0 !important;
        margin: 0 !important;
    }
    
    /* Hide upload-section divs that only contain whitespace */
    .upload-section:blank {
        display: none !important;
    }
    
    /* Additional file uploader enhancements */
    [data-testid="stFileUploader"] {
        background: rgba(255, 255, 255, 0.1) !important;
        backdrop-filter: blur(10px) !important;
        border-radius: 15px !important;
        border: 1px solid rgba(255, 255, 255, 0.2) !important;
        padding: 1rem !important;
    }
    
    [data-testid="stFileUploaderDropzone"] {
        background: rgba(255, 255, 255, 0.15) !important;
        backdrop-filter: blur(8px) !important;
        border: 1px solid rgba(255, 255, 255, 0.3) !important;
        border-radius: 12px !important;
        transition: all 0.3s ease !important;
        min-height: 120px !important;
    }
    
    [data-testid="stFileUploaderDropzone"]:hover {
        background: rgba(255, 255, 255, 0.25) !important;
        border: 1px solid rgba(255, 255, 255, 0.5) !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 25px rgba(0, 0, 0, 0.15) !important;
    }
    
    /* Widget labels with glass effect */
    [data-testid="stWidgetLabel"] {
        background: rgba(255, 255, 255, 0.1) !important;
        backdrop-filter: blur(5px) !important;
        border-radius: 8px !important;
        padding: 0.5rem !important;
        margin-bottom: 0.5rem !important;
    }
    
    /* Success state styling for uploaded files */
    .upload-success {
        background: linear-gradient(135deg, rgba(16, 185, 129, 0.2) 0%, rgba(16, 185, 129, 0.1) 100%) !important;
        border: 1px solid rgba(16, 185, 129, 0.4) !important;
        backdrop-filter: blur(15px) !important;
        border-radius: 15px !important;
        padding: 1rem !important;
        margin: 0.5rem 0 !important;
        box-shadow: 0 4px 20px rgba(16, 185, 129, 0.2) !important;
    }
    
    /* Hide file uploader when file is uploaded */
    .upload-section .stFileUploader {
        margin-bottom: 1rem;
    }
    
    /* Glass morphism styling for tabs */
    .stTabs [data-baseweb="tab-list"] {
        background: rgba(255, 255, 255, 0.15) !important;
        backdrop-filter: blur(20px) !important;
        -webkit-backdrop-filter: blur(20px) !important;
        border-radius: 20px !important;
        border: 1px solid rgba(255, 255, 255, 0.3) !important;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.1) !important;
        padding: 0.5rem !important;
        gap: 0.5rem !important;
    }
    
    .stTabs [data-baseweb="tab"] {
        background: rgba(255, 255, 255, 0.1) !important;
        color: #1e3c72 !important;
        border: 1px solid rgba(255, 255, 255, 0.2) !important;
        border-radius: 15px !important;
        padding: 0.75rem 1.5rem !important;
        font-weight: 500 !important;
        transition: all 0.3s ease !important;
        backdrop-filter: blur(10px) !important;
        -webkit-backdrop-filter: blur(10px) !important;
        box-shadow: 0 4px 16px rgba(0, 0, 0, 0.05) !important;
    }
    
    .stTabs [data-baseweb="tab"]:hover {
        background: rgba(255, 255, 255, 0.2) !important;
        border-color: rgba(255, 255, 255, 0.4) !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px rgba(0, 0, 0, 0.1) !important;
    }
    
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        background: rgba(255, 255, 255, 0.25) !important;
        color: #1e3c72 !important;
        border: 1px solid rgba(30, 60, 114, 0.4) !important;
        font-weight: 600 !important;
        box-shadow: 0 6px 24px rgba(30, 60, 114, 0.15) !important;
    }
    
    .stTabs [data-baseweb="tab-panel"] {
        background: rgba(255, 255, 255, 0.1) !important;
        backdrop-filter: blur(15px) !important;
        -webkit-backdrop-filter: blur(15px) !important;
        border-radius: 20px !important;
        border: 1px solid rgba(255, 255, 255, 0.2) !important;
        padding: 2rem !important;
        margin-top: 1rem !important;
        box-shadow: 0 8px 32px rgba(0, 0, 0, 0.1) !important;
    }
    
    /* Custom styling for Home buttons */
    [data-testid="stButton"]:has(button[aria-label=""]) button:contains("🏠 Home") {
        background: rgba(255, 255, 255, 0.9) !important;
        color: #1e3c72 !important;
        border: 1px solid rgba(30, 60, 114, 0.3) !important;
        border-radius: 20px !important;
        padding: 0.4rem 1rem !important;
        font-size: 0.9rem !important;
        font-weight: 500 !important;
        min-height: 36px !important;
        max-height: 36px !important;
        backdrop-filter: blur(10px) !important;
        transition: all 0.2s ease !important;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.1) !important;
    }
    
    /* Home button hover effect */
    [data-testid="stButton"]:has(button[aria-label=""]) button:contains("🏠 Home"):hover {
        background: rgba(255, 255, 255, 1) !important;
        border-color: #1e3c72 !important;
        box-shadow: 0 4px 12px rgba(30, 60, 114, 0.2) !important;
        transform: translateY(-1px) !important;
    }
    
    /* Simple home button styling - applied via JavaScript */
    .home-styled {
        background: rgba(255, 255, 255, 0.15) !important;
        color: #1e3c72 !important;
        border: 1px solid rgba(255, 255, 255, 0.4) !important;
        border-radius: 20px !important;
        padding: 0.4rem 1rem !important;
        font-size: 0.9rem !important;
        font-weight: 500 !important;
        min-height: 36px !important;
        max-height: 36px !important;
        backdrop-filter: blur(20px) !important;
        -webkit-backdrop-filter: blur(20px) !important;
        transition: all 0.2s ease !important;
        box-shadow: 0 4px 16px rgba(0, 0, 0, 0.1) !important;
        width: auto !important;
    }
    
    .home-styled:hover {
        background: rgba(255, 255, 255, 0.25) !important;
        border-color: rgba(255, 255, 255, 0.6) !important;
        box-shadow: 0 6px 20px rgba(30, 60, 114, 0.15) !important;
        transform: translateY(-2px) !important;
    }
</style>

<script>
function applyHomeButtonStyle() {
    // Find all buttons
    const buttons = document.querySelectorAll('button');
    
    buttons.forEach(button => {
        // Get all text content including nested elements
        const textContent = button.textContent || button.innerText || '';
        
        // Check if this is a Home button
        if (textContent.includes('🏠') && textContent.includes('Home')) {
            // Remove any existing styling classes first
            button.classList.remove('home-styled');
            
            // Add our custom class
            button.classList.add('home-styled');
            
            console.log('Applied home styling to button:', textContent.trim());
        }
    });
}

// Run the function on page load
document.addEventListener('DOMContentLoaded', applyHomeButtonStyle);

// Also run after a short delay to catch Streamlit's dynamic content
setTimeout(applyHomeButtonStyle, 100);
setTimeout(applyHomeButtonStyle, 500);
setTimeout(applyHomeButtonStyle, 1000);

// Set up a MutationObserver to catch new buttons
const observer = new MutationObserver(function(mutations) {
    let shouldRun = false;
    mutations.forEach(function(mutation) {
        if (mutation.addedNodes.length > 0) {
            shouldRun = true;
        }
    });
    
    if (shouldRun) {
        setTimeout(applyHomeButtonStyle, 50);
    }
});

// Start observing
observer.observe(document.body, {
    childList: true,
    subtree: true
});

// Re-apply on Streamlit page changes
window.addEventListener('focus', function() {
    setTimeout(applyHomeButtonStyle, 100);
});
</script>
</style>
""", unsafe_allow_html=True)

def initialize_session_state():
    """Initialize session state variables"""
    if 'comparison_results' not in st.session_state:
        st.session_state.comparison_results = None
    if 'output_file_data' not in st.session_state:
        st.session_state.output_file_data = None
    if 'comparison_complete' not in st.session_state:
        st.session_state.comparison_complete = False

def create_navigation_menu():
    """Create simplified navigation with home button in sidebar"""
    st.sidebar.markdown("### � Navigation")
    
    # Home button
    if st.sidebar.button("🏠 Back to Home", use_container_width=True):
        st.session_state.current_page = "Dashboard"
        st.rerun()
    
    # Show current page indicator
    if 'current_page' in st.session_state and st.session_state.current_page != "Dashboard":
        st.sidebar.markdown(f"**📍 Current Page:**")
        st.sidebar.markdown(f"🔸 {st.session_state.current_page}")
    
    # Navigation help
    st.sidebar.markdown("---")
    st.sidebar.markdown("### ℹ️ How to Navigate")
    st.sidebar.markdown("• **Dashboard:** Click tool cards to navigate")
    st.sidebar.markdown("• **Any Page:** Use 🏠 Home button to return")
    
    # Initialize current page if not set
    if 'current_page' not in st.session_state:
        st.session_state.current_page = "Dashboard"
    
    return st.session_state.current_page

def create_download_link(file_data, filename, link_text):
    """Create a download link for the results file"""
    b64 = base64.b64encode(file_data).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="{filename}" class="download-button">{link_text}</a>'
    return href

def validate_excel_file(uploaded_file, file_type):
    """Validate uploaded Excel file"""
    if uploaded_file is None:
        return False, f"Please upload {file_type} file"
    
    if not uploaded_file.name.endswith(('.xlsx', '.xls')):
        return False, f"{file_type} must be an Excel file (.xlsx or .xls)"
    
    try:
        # Try to read the file to validate it
        pd.read_excel(uploaded_file, nrows=1)
        return True, "File is valid"
    except Exception as e:
        return False, f"Error reading {file_type}: {str(e)}"

def display_file_info(uploaded_file, file_type):
    """Display information about uploaded file"""
    if uploaded_file:
        st.success(f"✅ {file_type} uploaded successfully!")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("📁 Filename", uploaded_file.name)
        with col2:
            st.metric("📊 File Size", f"{uploaded_file.size / 1024:.1f} KB")
        with col3:
            try:
                # Get basic info about the file
                df = pd.read_excel(uploaded_file, nrows=0)  # Just headers
                st.metric("📋 Columns", len(df.columns))
            except:
                st.metric("📋 Columns", "Unknown")

def run_comparison(ai_file, comparison_file):
    """Run the student data comparison"""
    
    # Create progress bar
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    try:
        status_text.text("🔧 Initializing comparator...")
        progress_bar.progress(10)
        
        if not COMPARATOR_AVAILABLE:
            st.error("❌ Student Data Comparator not available")
            return None, None, None
        
        # Initialize comparator with minimal logging for web app
        comparator = StudentDataComparator(log_level=logging.WARNING)
        
        status_text.text("📊 Loading AI extractor data...")
        progress_bar.progress(30)
        
        # Save uploaded files to temporary files
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_ai:
            tmp_ai.write(ai_file.getvalue())
            tmp_ai_path = tmp_ai.name
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_comp:
            tmp_comp.write(comparison_file.getvalue())
            tmp_comp_path = tmp_comp.name
        
        # Load data
        comparator.load_ai_extractor_data(tmp_ai_path, sheet_name="Extracted Data")
        
        status_text.text("📊 Loading comparison data...")
        progress_bar.progress(50)
        
        comparator.load_comparison_data(tmp_comp_path)
        
        status_text.text("🔍 Performing comparison...")
        progress_bar.progress(70)
        
        # Perform comparison
        results = comparator.compare_data(fuzzy_threshold=80)
        
        if 'error' in results:
            st.error(f"❌ Error during comparison: {results['error']}")
            return None, None
        
        status_text.text("📊 Generating results file...")
        progress_bar.progress(90)
        
        # Generate output file
        output_filename = f"student_comparison_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_output:
            tmp_output_path = tmp_output.name
        
        comparator.export_results(tmp_output_path)
        
        # Read the output file data
        with open(tmp_output_path, 'rb') as f:
            output_data = f.read()
        
        status_text.text("✅ Comparison complete!")
        progress_bar.progress(100)
        
        # Clean up temporary files
        os.unlink(tmp_ai_path)
        os.unlink(tmp_comp_path)
        os.unlink(tmp_output_path)
        
        return results, output_data, output_filename
        
    except Exception as e:
        st.error(f"❌ Error during comparison: {str(e)}")
        return None, None, None

def display_results(results):
    """Display comparison results in a beautiful format"""
    
    st.markdown('<div class="results-container">', unsafe_allow_html=True)
    
    # Main metrics
    st.markdown("## 📊 Comparison Results")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown(
            f'<div class="metric-card"><h3>{results["total_ai_students"]}</h3><p>AI Students</p></div>',
            unsafe_allow_html=True
        )
    
    with col2:
        st.markdown(
            f'<div class="metric-card"><h3>{results["total_comparison_students"]}</h3><p>Comparison Students</p></div>',
            unsafe_allow_html=True
        )
    
    with col3:
        st.markdown(
            f'<div class="metric-card"><h3>{results["matches_found"]}</h3><p>Matches Found</p></div>',
            unsafe_allow_html=True
        )
    
    with col4:
        st.markdown(
            f'<div class="metric-card"><h3>{results["match_rate"]:.1f}%</h3><p>Match Rate</p></div>',
            unsafe_allow_html=True
        )
    
    # Detailed breakdown
    st.markdown("### 📈 Detailed Breakdown")
    
    breakdown_col1, breakdown_col2 = st.columns(2)
    
    with breakdown_col1:
        if results['matches_found'] > 0:
            st.markdown(
                f'<div class="success-box">✅ <strong>{results["matches_found"]} students successfully matched</strong><br>'
                f'These students were found in both files and matched successfully.</div>',
                unsafe_allow_html=True
            )
    
    with breakdown_col2:
        unmatched_total = results['unmatched_ai'] + results['unmatched_comparison']
        if unmatched_total > 0:
            st.markdown(
                f'<div class="warning-box">⚠️ <strong>{unmatched_total} unmatched records</strong><br>'
                f'AI only: {results["unmatched_ai"]} | Comparison only: {results["unmatched_comparison"]}</div>',
                unsafe_allow_html=True
            )
    
    # Match rate visualization
    if results['match_rate'] >= 80:
        st.success(f"🎉 Excellent match rate of {results['match_rate']:.1f}%!")
    elif results['match_rate'] >= 60:
        st.warning(f"⚠️ Good match rate of {results['match_rate']:.1f}%, but some records were unmatched.")
    else:
        st.error(f"❌ Low match rate of {results['match_rate']:.1f}%. You may need to check data quality or adjust settings.")
    
    # Suggestions based on results
    if results['unmatched_ai'] > 0:
        st.info(f"💡 **{results['unmatched_ai']} students from AI data were not matched**: These might be new students not in the comparison file.")
    
    if results['unmatched_comparison'] > 0:
        st.info(f"💡 **{results['unmatched_comparison']} students from comparison data were not matched**: These might be students not processed by the AI extractor.")
    
    st.markdown('</div>', unsafe_allow_html=True)

def show_dashboard():
    """Dashboard page with overview and quick access"""
    
    # Quick access navigation cards
    st.markdown("## Select an individual process or do all of them at once ⚡")
    st.markdown("Click any card below to access that tool:")
    
    # Add spacing
    st.markdown("<div style='margin-bottom: 1rem;'></div>", unsafe_allow_html=True)
    
    # First row of buttons
    col1, col2 = st.columns(2, gap="medium")
    
    with col1:
        if st.button("📄  Word to PDF Conversion", key="nav_word_to_pdf", use_container_width=True):
            st.session_state.current_page = "Word to PDF"
            st.rerun()
    
    with col2:
        if st.button("🤖  AI Data Extraction", key="nav_ai_extraction", use_container_width=True):
            st.session_state.current_page = "PDF Extraction"
            st.rerun()
    
    # Add spacing between rows
    st.markdown("<div style='margin: 1rem 0;'></div>", unsafe_allow_html=True)
    
    # Second row of buttons
    col3, col4 = st.columns(2, gap="medium")
    
    with col3:
        if st.button("🔍  Data Validation", key="nav_data_validation", use_container_width=True):
            st.session_state.current_page = "Data Validation"
            st.rerun()
    
    with col4:
        if st.button("🚌  Traversa Preparation", key="nav_traversa_prep", use_container_width=True):
            st.session_state.current_page = "Traversa Preparation"
            st.rerun()
    
    # Add spacing before all-in-one button
    st.markdown("<div style='margin: 1.5rem 0;'></div>", unsafe_allow_html=True)
    
    # Full-width all-in-one button
    if st.button("⚡  All-in-One Processing", key="nav_all_in_one", use_container_width=True):
        st.session_state.current_page = "All-in-One"
        st.rerun()
    
    st.markdown("## Welcome to Voigts Student Opt-In Data Management Suite")
    
    st.markdown("""
    **Voigt's Bus Companies** presents this comprehensive student data processing platform, designed for student transportation opt-in data:
    
    ### 🚌 Transportation-Focused Tools:
    - **Word to PDF Converter** - Convert student transportation forms from Word documents to PDF format
    - **AI Data Extraction** - Use advanced AI to extract transportation preferences and student data
    - **Data Validation** - Compare and validate student information against district databases
    - **Traversa Preparation** - Format data specifically for Traversa routing software upload with transportation analysis
    - **All-in-One Processor** - Complete automated pipeline from forms to route-ready data
    
    ### 🍎 Designed for School Districts:
    1. **Streamlined Workflow** - From paper forms to digital routing in minutes
    2. **Transportation Intelligence** - Automatic categorization of AM/PM transportation needs
    3. **Color-Coded Results** - Easy identification of changes and transportation requirements
    4. **Multi-District Ready** - Works with different data formats and requirements
    3. Download your processed results
    
    Choose a tool from the sidebar to begin!
    """)

def convert_docx_to_pdf_silent(docx_file, output_dir):
    """Convert a DOCX file to PDF using python-docx and reportlab (silent, no Word app needed)"""
    try:
        # Read the DOCX document
        docx_data = BytesIO(docx_file.read())
        document = Document(docx_data)
        
        # Generate output PDF path
        pdf_filename = docx_file.name.replace('.docx', '.pdf').replace('.DOCX', '.pdf')
        output_pdf_path = os.path.join(output_dir, pdf_filename)
        
        # Create PDF using ReportLab
        pdf_buffer = BytesIO()
        doc = SimpleDocTemplate(pdf_buffer, pagesize=letter, 
                              rightMargin=72, leftMargin=72, 
                              topMargin=72, bottomMargin=18)
        
        # Get styles
        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        heading_style = styles['Heading1']
        
        # Story array to hold document content
        story = []
        
        # Process each paragraph in the Word document
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                # Determine if it's a heading based on paragraph style
                if paragraph.style.name.startswith('Heading'):
                    story.append(Paragraph(paragraph.text, heading_style))
                else:
                    story.append(Paragraph(paragraph.text, normal_style))
                story.append(Spacer(1, 12))
        
        # Process tables if any
        for table in document.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            
            if table_data:
                pdf_table = Table(table_data)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(pdf_table)
                story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        
        # Save to file
        with open(output_pdf_path, 'wb') as f:
            f.write(pdf_buffer.getvalue())
        
        return output_pdf_path, pdf_filename
        
    except Exception as e:
        raise Exception(f"Silent conversion failed: {str(e)}")

def convert_docx_to_pdf_fallback(docx_file, output_dir):
    """Fallback conversion using docx2pdf (may open Word app)"""
    try:
        # Create temporary file for the docx
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_docx:
            temp_docx.write(docx_file.read())
            temp_docx_path = temp_docx.name
        
        # Generate output PDF path
        pdf_filename = docx_file.name.replace('.docx', '.pdf').replace('.DOCX', '.pdf')
        output_pdf_path = os.path.join(output_dir, pdf_filename)
        
        # Convert using docx2pdf
        docx_convert(temp_docx_path, output_pdf_path)
        
        # Clean up temporary file
        os.unlink(temp_docx_path)
        
        return output_pdf_path, pdf_filename
        
    except Exception as e:
        # Clean up on error
        if 'temp_docx_path' in locals():
            try:
                os.unlink(temp_docx_path)
            except:
                pass
        raise e

def convert_docx_to_pdf(docx_file, output_dir):
    """Convert DOCX to PDF with automatic method selection"""
    # Try silent conversion first (preferred - no Word app needed)
    if DOCX_TO_PDF_AVAILABLE:
        try:
            return convert_docx_to_pdf_silent(docx_file, output_dir)
        except Exception as e:
            st.warning(f"⚠️ Silent conversion failed: {str(e)}")
            if DOCX2PDF_FALLBACK:
                st.info("🔄 Trying fallback method (may open Word app)...")
                return convert_docx_to_pdf_fallback(docx_file, output_dir)
            else:
                raise e
    # Use fallback method if silent not available
    elif DOCX2PDF_FALLBACK:
        st.info("ℹ️ Using system Word application for conversion...")
        return convert_docx_to_pdf_fallback(docx_file, output_dir)
    else:
        raise Exception("No conversion method available. Please install python-docx and reportlab, or docx2pdf.")
    """Create a ZIP file containing multiple PDFs"""
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for file_path in file_paths:
            if os.path.exists(file_path):
                zip_file.write(file_path, os.path.basename(file_path))
    
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

def create_download_zip(file_paths, zip_name):
    """Create a ZIP file containing multiple PDFs"""
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for file_path in file_paths:
            if os.path.exists(file_path):
                zip_file.write(file_path, os.path.basename(file_path))
    
    zip_buffer.seek(0)
    return zip_buffer.getvalue()
    """Word to PDF conversion page with actual functionality"""
    st.markdown('<h1 class="main-header"><span class="emoji-icon">📄</span>Word to PDF Converter</h1>', unsafe_allow_html=True)
    st.markdown('<p class="subtitle">Convert Word documents to PDF format for processing</p>', unsafe_allow_html=True)
    
    # Check if docx conversion is available
    if not DOCX_TO_PDF_AVAILABLE and not DOCX2PDF_FALLBACK:
        st.error("❌ **Word to PDF conversion not available**")
        st.markdown("""
        To use this feature, please install the required dependencies:
        ```bash
        pip install python-docx reportlab
        ```
        Or alternatively:
        ```bash
        pip install docx2pdf
        ```
        """)
        return
    
    # Display conversion method being used
    if DOCX_TO_PDF_AVAILABLE:
        st.success("✅ **Silent conversion enabled** - No Word application required!")
    elif DOCX2PDF_FALLBACK:
        st.warning("⚠️ **System Word conversion** - May require permission prompts")

def show_word_to_pdf():
    """Word to PDF conversion page with actual functionality"""
    # Home button at the top
    col1, col2, col3 = st.columns([1, 8, 1])
    with col1:
        if st.button("🏠 Home", key="home_word_pdf"):
            st.session_state.current_page = "Dashboard"
            st.rerun()
    
    st.markdown('<h1 class="main-header"><span class="emoji-icon">📄</span>Voigts Word to PDF Converter</h1>', unsafe_allow_html=True)
    st.markdown('<p class="subtitle">Convert transportation forms from Word documents to PDF format for AI processing</p>', unsafe_allow_html=True)
    
    # Conversion mode selection
    st.markdown("## ⚙️ Conversion Mode")
    conversion_mode = st.radio(
        "Choose conversion mode:",
        ["Single File", "Multiple Files"],
        key="conversion_mode"
    )
    
    if conversion_mode == "Single File":
        st.markdown("### 📄 Single File Conversion")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown('<div class="upload-section">', unsafe_allow_html=True)
            st.markdown("### 📁 Upload Word Document")
            
            uploaded_file = st.file_uploader(
                "Choose a Word document",
                type=['docx', 'DOCX'],
                key="single_word_file",
                help="Upload a .docx file to convert to PDF"
            )
            
            if uploaded_file:
                st.markdown('<div class="upload-success">', unsafe_allow_html=True)
                st.markdown(f"### ✅ **{uploaded_file.name}**")
                st.markdown("📄 *Ready for conversion*")
                st.markdown('</div>', unsafe_allow_html=True)
                
                # File info
                st.markdown("**File Details:**")
                st.markdown(f"- **Size:** {uploaded_file.size / 1024:.1f} KB")
                st.markdown(f"- **Type:** {uploaded_file.type}")
            
            st.markdown('</div>', unsafe_allow_html=True)
        
        with col2:
            st.markdown('<div class="upload-section">', unsafe_allow_html=True)
            st.markdown("### 🎯 Conversion Options")
            
            if uploaded_file:
                st.success("✅ Ready to convert!")
                
                if st.button("🚀 Convert to PDF", type="primary"):
                    try:
                        with st.spinner("🔄 Converting document..."):
                            # Create temporary directory for output
                            with tempfile.TemporaryDirectory() as temp_dir:
                                pdf_path, pdf_filename = convert_docx_to_pdf(uploaded_file, temp_dir)
                                
                                # Read the PDF file
                                with open(pdf_path, 'rb') as pdf_file:
                                    pdf_data = pdf_file.read()
                                
                                st.success(f"✅ Successfully converted to {pdf_filename}")
                                
                                # Download button
                                st.download_button(
                                    label="📥 Download PDF",
                                    data=pdf_data,
                                    file_name=pdf_filename,
                                    mime="application/pdf"
                                )
                                
                    except Exception as e:
                        st.error(f"❌ Conversion failed: {str(e)}")
            else:
                st.info("📤 Upload a Word document to begin conversion")
            
            st.markdown('</div>', unsafe_allow_html=True)
    
    else:  # Multiple Files
        st.markdown("### 📁 Batch File Conversion")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown('<div class="upload-section">', unsafe_allow_html=True)
            st.markdown("### 📁 Upload Word Documents")
            
            uploaded_files = st.file_uploader(
                "Choose Word documents",
                type=['docx', 'DOCX'],
                accept_multiple_files=True,
                key="multiple_word_files",
                help="Upload multiple .docx files to convert to PDF"
            )
            
            if uploaded_files:
                st.markdown('<div class="upload-success">', unsafe_allow_html=True)
                st.markdown(f"### ✅ **{len(uploaded_files)} files uploaded**")
                st.markdown("📄 *Ready for batch conversion*")
                st.markdown('</div>', unsafe_allow_html=True)
                
                # Show file list
                st.markdown("**Files to convert:**")
                for i, file in enumerate(uploaded_files, 1):
                    st.markdown(f"{i}. {file.name} ({file.size / 1024:.1f} KB)")
            
            st.markdown('</div>', unsafe_allow_html=True)
        
        with col2:
            st.markdown('<div class="upload-section">', unsafe_allow_html=True)
            st.markdown("### 🎯 Batch Conversion")
            
            if uploaded_files:
                st.success(f"✅ Ready to convert {len(uploaded_files)} files!")
                
                if st.button("🚀 Convert All to PDF", type="primary"):
                    try:
                        with st.spinner(f"🔄 Converting {len(uploaded_files)} documents..."):
                            # Create temporary directory for output
                            with tempfile.TemporaryDirectory() as temp_dir:
                                pdf_paths = []
                                
                                # Convert each file
                                progress_bar = st.progress(0)
                                for i, file in enumerate(uploaded_files):
                                    try:
                                        pdf_path, pdf_filename = convert_docx_to_pdf(file, temp_dir)
                                        pdf_paths.append(pdf_path)
                                        st.write(f"✅ Converted: {file.name} → {pdf_filename}")
                                    except Exception as e:
                                        st.error(f"❌ Failed to convert {file.name}: {str(e)}")
                                    
                                    # Update progress
                                    progress_bar.progress((i + 1) / len(uploaded_files))
                                
                                if pdf_paths:
                                    # Create ZIP file with all PDFs
                                    zip_data = create_download_zip(pdf_paths, "converted_pdfs.zip")
                                    
                                    st.success(f"✅ Successfully converted {len(pdf_paths)} documents!")
                                    
                                    # Download ZIP button
                                    st.download_button(
                                        label="📥 Download All PDFs (ZIP)",
                                        data=zip_data,
                                        file_name="converted_pdfs.zip",
                                        mime="application/zip"
                                    )
                                else:
                                    st.error("❌ No files were successfully converted")
                                    
                    except Exception as e:
                        st.error(f"❌ Batch conversion failed: {str(e)}")
            else:
                st.info("📤 Upload Word documents to begin batch conversion")
            
            st.markdown('</div>', unsafe_allow_html=True)
    
    # Usage instructions
    st.markdown("---")
    st.markdown("## 📋 Instructions")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        ### ✅ Supported Features:
        - **.docx file conversion** to PDF
        - **Single file** conversion with immediate download
        - **Batch processing** with ZIP download
        - **File validation** and error handling
        - **Progress tracking** for batch operations
        """)
    
    with col2:
        st.markdown("""
        ### 📝 Usage Tips:
        - Upload **.docx files only** (Word 2007+)
        - **File size limit:** Recommended under 10MB per file
        - **Batch mode:** Process up to 50 files at once
        - **Output:** High-quality PDF files ready for AI extraction
        """)

def extract_data_from_pdfs(pdf_files, progress_callback=None, model_id="auto", extract_options=None, file_models=None):
    """Extract data from uploaded PDF files using Azure AI"""
    try:
        # Get Azure credentials from secrets
        try:
            endpoint = st.secrets["azure"]["endpoint"]
            api_key = st.secrets["azure"]["api_key"]
        except:
            return {"error": "Azure credentials not configured"}
        
        # Initialize Azure AI client
        client = DocumentAnalysisClient(
            endpoint=endpoint,
            credential=AzureKeyCredential(api_key)
        )
        
        extracted_data = []
        
        for i, pdf_file in enumerate(pdf_files):
            if progress_callback:
                progress_callback(i / len(pdf_files))
            
            # Reset file pointer
            pdf_file.seek(0)
            
            # Analyze document with Azure AI
            poller = client.begin_analyze_document(
                "prebuilt-document", 
                document=pdf_file.read()
            )
            result = poller.result()
            
            # Extract text content
            file_data = {
                'filename': pdf_file.name,
                'extracted_text': [],
                'tables': [],
                'key_value_pairs': []
            }
            
            # Extract text from pages
            for page in result.pages:
                for line in page.lines:
                    file_data['extracted_text'].append(line.content)
            
            # Extract tables if available
            for table in result.tables:
                table_data = []
                for row in table.cells:
                    table_data.append({
                        'row_index': row.row_index,
                        'column_index': row.column_index,
                        'content': row.content
                    })
                file_data['tables'].append(table_data)
            
            # Extract key-value pairs
            for kv_pair in result.key_value_pairs:
                if kv_pair.key and kv_pair.value:
                    file_data['key_value_pairs'].append({
                        'key': kv_pair.key.content,
                        'value': kv_pair.value.content
                    })
            
            extracted_data.append(file_data)
        
        if progress_callback:
            progress_callback(1.0)
        
        return {"success": True, "data": extracted_data, "excel_data": None, "filename": f"extracted_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"}
        
    except Exception as e:
        return {"error": str(e)}

def show_pdf_extraction():
    """PDF data extraction page with AI functionality"""
    # Home button at the top
    col1, col2, col3 = st.columns([1, 8, 1])
    with col1:
        if st.button("🏠 Home", key="home_pdf_extract"):
            st.session_state.current_page = "Dashboard"
            st.rerun()
    
    st.markdown('<h1 class="main-header"><span class="emoji-icon">🤖</span>Voigts AI Data Extraction</h1>', unsafe_allow_html=True)
    st.markdown('<p class="subtitle">Extract transportation preferences and student data using advanced AI technology</p>', unsafe_allow_html=True)
    
    # Check if AI extractor is available
    if not AI_EXTRACTOR_AVAILABLE:
        st.error("❌ **AI PDF Extractor not available**")
        st.markdown("""
        To use this feature, please ensure the following dependencies are installed:
        ```bash
        pip install azure-ai-formrecognizer azure-core pandas openpyxl
        ```
        **For local use:** Configure Azure credentials in `config.ini`  
        **For cloud deployment:** Set Azure credentials in Streamlit Cloud secrets
        """)
        return
    
    # Configuration section
    st.markdown("## ⚙️ AI Configuration")
    
    # Model Management Section
    st.markdown("### 🧠 AI Model Management")
    
    # Create tabs for model selection and custom model upload
    model_tab1, model_tab2 = st.tabs(["📋 Select Model", "🔧 Manage Custom Models"])
    
    with model_tab1:
        config_col1, config_col2 = st.columns(2)
        
        with config_col1:
            st.markdown("#### Available Models")
            
            # Get available models (both built-in and custom)
            available_models = {
                "Auto-detect (Recommended)": "auto",
                "ROCORI Transportation Forms": "rocorioptin", 
                "General Document": "prebuilt-document",
                "Daycare Forms": "daycareoptin2"
            }
            
            # Check for custom models in session state
            if 'custom_models' in st.session_state:
                for name, model_id in st.session_state.custom_models.items():
                    available_models[f"Custom: {name}"] = model_id
            
            model_type = st.selectbox(
                "Choose AI model:",
                list(available_models.keys()),
                help="Select the AI model that best matches your document type"
            )
            
            selected_model_id = available_models[model_type]
            
            # Show model info
            if model_type == "Auto-detect (Recommended)":
                st.info("🤖 Will automatically select the best model based on document content")
            elif "Custom:" in model_type:
                st.success(f"🎯 Using your custom model: {selected_model_id}")
            else:
                st.info(f"📋 Using model: {selected_model_id}")
            
        with config_col2:
            st.markdown("#### Extraction Options") 
            extract_tables = st.checkbox("Extract Tables", value=True)
            extract_forms = st.checkbox("Extract Form Fields", value=True)
            extract_text = st.checkbox("Include Full Text", value=False)
            confidence_threshold = st.slider("Confidence Threshold", 0.5, 1.0, 0.7, 0.05,
                                            help="Minimum confidence level for extracted data")
    
    with model_tab2:
        st.markdown("#### Add Custom AI Models")
        st.markdown("Upload your own trained Azure Document Intelligence models for specialized document types.")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Add New Model:**")
            custom_model_name = st.text_input("Model Name", 
                                            placeholder="e.g., My School Forms",
                                            help="Give your model a descriptive name")
            custom_model_id = st.text_input("Model ID", 
                                          placeholder="e.g., myschoolforms-v1",
                                          help="Your Azure Document Intelligence model ID")
            
            if st.button("➕ Add Custom Model"):
                if custom_model_name and custom_model_id:
                    if 'custom_models' not in st.session_state:
                        st.session_state.custom_models = {}
                    
                    st.session_state.custom_models[custom_model_name] = custom_model_id
                    st.success(f"✅ Added custom model: {custom_model_name}")
                    st.rerun()
                else:
                    st.error("Please provide both model name and ID")
        
        with col2:
            st.markdown("**Current Custom Models:**")
            if 'custom_models' in st.session_state and st.session_state.custom_models:
                for name, model_id in st.session_state.custom_models.items():
                    col_name, col_id, col_remove = st.columns([3, 3, 1])
                    with col_name:
                        st.text(name)
                    with col_id:
                        st.code(model_id, language=None)
                    with col_remove:
                        if st.button("🗑️", key=f"remove_{name}", help="Remove this model"):
                            del st.session_state.custom_models[name]
                            st.rerun()
            else:
                st.info("No custom models added yet")
                
        st.markdown("---")
        st.markdown("**ℹ️ How to add custom models:**")
        st.markdown("""
        1. **Train a model** in Azure Document Intelligence Studio
        2. **Copy the model ID** from your Azure resource
        3. **Add it here** with a descriptive name
        4. **Use it** for specialized document extraction
        """)
        
        st.markdown("**📚 Supported model types:**")
        st.markdown("- Custom document models")
        st.markdown("- Composed models") 
        st.markdown("- Custom classification models")
        st.markdown("- Custom extraction models")
        
        # Import/Export Models
        st.markdown("---")
        st.markdown("#### 📥📤 Import/Export Models")
        
        import_col, export_col = st.columns(2)
        
        with import_col:
            st.markdown("**Import Models:**")
            
            # JSON file upload
            uploaded_models_file = st.file_uploader(
                "Upload models JSON file",
                type=['json'],
                help="Upload a JSON file containing model definitions"
            )
            
            if uploaded_models_file:
                try:
                    import json
                    models_data = json.load(uploaded_models_file)
                    
                    if st.button("📥 Import Models"):
                        if 'custom_models' not in st.session_state:
                            st.session_state.custom_models = {}
                        
                        imported_count = 0
                        for name, model_id in models_data.items():
                            st.session_state.custom_models[name] = model_id
                            imported_count += 1
                        
                        st.success(f"✅ Imported {imported_count} models")
                        st.rerun()
                        
                except Exception as e:
                    st.error(f"Error importing models: {str(e)}")
            
            # Manual JSON input
            st.markdown("**Or paste JSON:**")
            models_json = st.text_area(
                "Models JSON",
                placeholder='{"Model Name": "model-id", "Another Model": "another-id"}',
                help="Paste JSON with model name-ID pairs"
            )
            
            if models_json and st.button("📥 Import from JSON"):
                try:
                    import json
                    models_data = json.loads(models_json)
                    
                    if 'custom_models' not in st.session_state:
                        st.session_state.custom_models = {}
                    
                    imported_count = 0
                    for name, model_id in models_data.items():
                        st.session_state.custom_models[name] = model_id
                        imported_count += 1
                    
                    st.success(f"✅ Imported {imported_count} models")
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Error parsing JSON: {str(e)}")
        
        with export_col:
            st.markdown("**Export Models:**")
            
            if 'custom_models' in st.session_state and st.session_state.custom_models:
                import json
                models_json = json.dumps(st.session_state.custom_models, indent=2)
                
                st.download_button(
                    label="📤 Download Models JSON",
                    data=models_json,
                    file_name=f"custom_models_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json",
                    help="Download your custom models as a JSON file"
                )
                
                st.markdown("**Preview:**")
                st.code(models_json, language="json")
            else:
                st.info("No custom models to export")
    
    # Upload section
    st.markdown("## 📄 Upload PDF Files")
    
    # Show current model status
    st.markdown("### 🧠 Current AI Model Configuration")
    status_col1, status_col2, status_col3 = st.columns(3)
    
    with status_col1:
        st.markdown(f"**🤖 Model:** `{model_type}`")
    with status_col2:
        st.markdown(f"**🆔 ID:** `{selected_model_id}`")
    with status_col3:
        if selected_model_id == "auto":
            st.markdown("**🎯 Mode:** Auto-detect")
        elif "Custom:" in model_type:
            st.markdown("**🎯 Mode:** Custom Model")
        else:
            st.markdown("**🎯 Mode:** Built-in Model")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown('<div class="upload-section">', unsafe_allow_html=True)
        st.markdown("### 📄 Upload PDF Documents")
        
        uploaded_pdfs = st.file_uploader(
            "Choose PDF files",
            type=['pdf'],
            accept_multiple_files=True,
            key="pdf_extraction_files",
            help="Upload PDF files to extract data from"
        )
        
        if uploaded_pdfs:
            st.markdown('<div class="upload-success">', unsafe_allow_html=True)
            st.markdown(f"### ✅ **{len(uploaded_pdfs)} files uploaded**")
            st.markdown("🤖 *Ready for AI extraction*")
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Show file list
            st.markdown("**Files to process:**")
            total_size = 0
            for i, file in enumerate(uploaded_pdfs, 1):
                size_kb = file.size / 1024
                total_size += size_kb
                st.markdown(f"{i}. {file.name} ({size_kb:.1f} KB)")
            
            st.markdown(f"**Total size:** {total_size:.1f} KB")
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="upload-section">', unsafe_allow_html=True)
        st.markdown("### 🎯 AI Extraction")
        
        if uploaded_pdfs:
            st.success(f"✅ Ready to process {len(uploaded_pdfs)} files!")
            
            # Estimation
            estimated_time = len(uploaded_pdfs) * 5  # Rough estimate: 5 seconds per PDF
            st.info(f"⏱️ Estimated processing time: ~{estimated_time} seconds")
            
            # Advanced batch processing options
            with st.expander("🔧 Advanced Batch Options"):
                st.markdown("#### 📋 Per-File Model Selection")
                st.markdown("Override the default model for specific files:")
                
                file_models = {}
                for i, file in enumerate(uploaded_pdfs):
                    col_file, col_model = st.columns([2, 1])
                    with col_file:
                        st.markdown(f"📄 {file.name}")
                    with col_model:
                        file_model = st.selectbox(
                            "Model",
                            ["Use Default"] + list(available_models.keys()),
                            key=f"file_model_{i}",
                            label_visibility="collapsed"
                        )
                        if file_model != "Use Default":
                            file_models[file.name] = available_models[file_model]
                
                if file_models:
                    st.success(f"✅ Custom models set for {len(file_models)} files")
                
                # Save file models to session state
                st.session_state.file_models = file_models
            
            if st.button("🚀 Start AI Extraction", type="primary"):
                try:
                    # Progress tracking
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    def update_progress(message, progress):
                        progress_bar.progress(progress)
                        status_text.text(f"🤖 {message}")
                    
                    with st.spinner("🤖 AI is analyzing your documents..."):
                        # Prepare extraction options
                        extract_options = {
                            'tables': extract_tables,
                            'forms': extract_forms,
                            'text': extract_text,
                            'confidence_threshold': confidence_threshold
                        }
                        
                        # Get file-specific models if any
                        file_models = st.session_state.get('file_models', {})
                        
                        excel_data, filename, extracted_data = extract_data_from_pdfs(
                            uploaded_pdfs, 
                            progress_callback=update_progress,
                            model_id=selected_model_id,
                            extract_options=extract_options,
                            file_models=file_models
                        )
                    
                    # Clear progress indicators
                    progress_bar.empty()
                    status_text.empty()
                    
                    st.success(f"✅ Successfully extracted data from {len(uploaded_pdfs)} files!")
                    
                    # Store results in session state
                    st.session_state.extraction_results = extracted_data
                    st.session_state.extraction_excel_data = excel_data
                    st.session_state.extraction_filename = filename
                    st.session_state.extraction_complete = True
                    
                except Exception as e:
                    st.error(f"❌ Extraction failed: {str(e)}")
                    st.markdown("Please check your Azure credentials and try again.")
        else:
            st.info("📤 Upload PDF files to begin AI extraction")
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Results section
    if hasattr(st.session_state, 'extraction_complete') and st.session_state.extraction_complete:
        st.markdown("---")
        st.markdown("## 📊 Extraction Results")
        
        if hasattr(st.session_state, 'extraction_results'):
            results = st.session_state.extraction_results
            
            # Summary metrics
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric("Files Processed", len(results))
            with col2:
                successful = len([r for r in results if 'error' not in r])
                st.metric("Successful", successful)
            with col3:
                errors = len([r for r in results if 'error' in r])
                st.metric("Errors", errors)
            with col4:
                avg_confidence = sum([r.get('confidence', 0) for r in results if 'confidence' in r]) / len(results) if results else 0
                st.metric("Avg Confidence", f"{avg_confidence:.1%}")
            
            # Detailed results
            st.markdown("### 📋 Detailed Results")
            
            for result in results:
                # Create a more informative title with model info
                title_parts = [f"📄 {result.get('source_file', 'Unknown File')}"]
                if 'model_used' in result:
                    title_parts.append(f"🤖 Model: {result['model_used']}")
                
                with st.expander(" | ".join(title_parts)):
                    if 'error' in result:
                        st.error(f"❌ Error: {result['error']}")
                        if 'model_used' in result:
                            st.markdown(f"**🤖 Model used:** `{result['model_used']}`")
                    else:
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.markdown("**Document Info:**")
                            st.markdown(f"- **Type:** {result.get('document_type', 'Unknown')}")
                            st.markdown(f"- **Confidence:** {result.get('confidence', 0):.1%}")
                            st.markdown(f"- **Pages:** {result.get('page_count', 0)}")
                            if 'model_used' in result:
                                st.markdown(f"- **🤖 AI Model:** `{result['model_used']}`")
                        
                        with col2:
                            st.markdown("**Extracted Fields:**")
                            field_count = 0
                            for key, value in result.items():
                                if key not in ['source_file', 'extraction_timestamp', 'document_type', 'confidence', 'page_count', 'tables', 'full_text', 'model_used']:
                                    if value and str(value).strip():
                                        st.markdown(f"- **{key}:** {value}")
                                        field_count += 1
                            
                            if field_count == 0:
                                st.markdown("*No form fields extracted*")
            
            # Download section
            st.markdown("### 💾 Download Results")
            
            if hasattr(st.session_state, 'extraction_excel_data'):
                download_col1, download_col2 = st.columns([2, 1])
                
                with download_col1:
                    st.markdown(f"📁 **{st.session_state.extraction_filename}**")
                    st.markdown("📊 Complete extraction results with all extracted data organized in sheets")
                    st.markdown("- **Summary Sheet:** Overview and statistics")
                    st.markdown("- **Extracted Data Sheet:** All extracted fields by file")
                
                with download_col2:
                    st.download_button(
                        label="� Download Excel Results",
                        data=st.session_state.extraction_excel_data,
                        file_name=st.session_state.extraction_filename,
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
            
            # Reset button
            if st.button("🔄 Process New Files"):
                # Clear session state
                for key in ['extraction_results', 'extraction_excel_data', 'extraction_filename', 'extraction_complete']:
                    if hasattr(st.session_state, key):
                        delattr(st.session_state, key)
                st.rerun()
    
    # Usage instructions
    st.markdown("---")
    st.markdown("## 📋 How It Works")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
        ### 🤖 AI Technology:
        - **Azure Document Intelligence** for form recognition
        - **Custom trained models** for ROCORI transportation forms
        - **Smart model selection** based on document content
        - **High accuracy** extraction with confidence scores
        - **Rate-limited processing** to respect API limits
        """)
    
    with col2:
        st.markdown("""
        ### 📄 Supported Documents:
        - **Student transportation forms**
        - **Opt-in/opt-out forms**
        - **Application forms**
        - **Structured documents with fields**
        - **Tables and form data**
        """)

def show_data_validation():
    """Enhanced data validation page with multiple comparison modes"""
    # Home button at the top
    col1, col2, col3 = st.columns([1, 8, 1])
    with col1:
        if st.button("🏠 Home", key="home_data_validation"):
            st.session_state.current_page = "Dashboard"
            st.rerun()
    
    st.markdown('<h1 class="main-header"><span class="emoji-icon">🔍</span>Voigts Data Validation System</h1>', unsafe_allow_html=True)
    st.markdown('<p class="subtitle">Compare and validate student transportation data against district databases with precision matching</p>', unsafe_allow_html=True)
    
    # Create tabs for different validation modes
    ai_tab, general_tab = st.tabs(["🤖 AI Data Validation", "📊 General Excel Comparison"])
    
    with ai_tab:
        st.markdown("### 🤖 AI Extracted Data Validation")
        st.markdown("**Primary Feature:** Compare AI-extracted student data against existing records")
        
        # File upload section for AI validation
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown('<div class="upload-section">', unsafe_allow_html=True)
            
            # Always show the header
            st.markdown("### 🧑‍🎓 Processed Student Opt-In Forms")
            st.markdown("📄 Upload the Excel output from AI Data Extraction - Select General Excel Comparison above to use any Excel file")
            
            ai_file = st.file_uploader(
                "Choose AI extractor file",
                type=['xlsx', 'xls'],
                key="ai_file",
                help="🤖Upload the Excel output from the previous step 'AI Opt-Form Data Extraction'",
                label_visibility="visible"
            )
            
            if ai_file:
                # File uploaded - show success state instead of description
                st.markdown('<div class="upload-success">', unsafe_allow_html=True)
                st.markdown(f"### ✅ **{ai_file.name}**")
                st.markdown("🤖 *Ready for processing*")
                st.markdown('</div>', unsafe_allow_html=True)
            
            st.markdown('</div>', unsafe_allow_html=True)
        
        with col2:
            st.markdown('<div class="upload-section">', unsafe_allow_html=True)
            
            # Always show the header
            st.markdown("### 📋 Traversa or Skyward Current Student Export")
            st.markdown("Upload your current student database export for comparison")
            
            comparison_file = st.file_uploader(
                "Choose comparison file",
                type=['xlsx', 'xls'],
                key="comparison_file",
                help="Upload your current student database export for comparison and validation",
                label_visibility="visible"
            )
            
            if comparison_file:
                # File uploaded - show success state instead of description
                st.markdown('<div class="upload-success">', unsafe_allow_html=True)
                st.markdown(f"### ✅ **{comparison_file.name}**")
                st.markdown("📋 *Ready for comparison*")
                st.markdown('</div>', unsafe_allow_html=True)
            
            st.markdown('</div>', unsafe_allow_html=True)
        
        # AI-specific configuration
        st.markdown("## ⚙️ AI Validation Settings")
        
        settings_col1, settings_col2, settings_col3 = st.columns(3)
        
        with settings_col1:
            ai_fuzzy_threshold = st.slider(
                "🎯 Matching Sensitivity",
                min_value=50,
                max_value=100,
                value=80,
                step=5,
                help="Higher values = stricter matching (fewer false positives), Lower values = more lenient matching (more matches)",
                key="ai_fuzzy_threshold"
            )
        
        with settings_col2:
            ai_max_results = st.number_input(
                "📊 Max Results to Display",
                min_value=10,
                max_value=1000,
                value=100,
                step=10,
                help="Limit the number of comparison results to display",
                key="ai_max_results"
            )
        
        with settings_col3:
            ai_show_preview = st.checkbox(
                "👀 Show File Preview", 
                value=True,
                help="Display a preview of the uploaded files",
                key="ai_show_preview"
            )
        
        # Processing section for AI validation
        if ai_file and comparison_file:
            if st.button("🚀 Start AI Data Validation", type="primary", key="start_ai_validation"):
                try:
                    with st.spinner("🤖 Validating AI extracted data..."):
                        results, output_data, output_filename = process_comparison(
                            ai_file, comparison_file, ai_fuzzy_threshold, ai_max_results
                        )
                    
                    if results and output_data:
                        st.session_state.ai_validation_results = results
                        st.session_state.ai_validation_output = output_data
                        st.session_state.ai_validation_filename = output_filename
                        st.success("✅ AI data validation completed!")
                        st.rerun()
                    else:
                        st.error("❌ Validation failed. Please check your files and try again.")
                        
                except Exception as e:
                    st.error(f"❌ Error during validation: {str(e)}")
        else:
            st.info("📤 Upload both AI extracted data and comparison data to begin validation")
        
        # Display AI validation results
        if hasattr(st.session_state, 'ai_validation_results') and st.session_state.ai_validation_results:
            display_validation_results(st.session_state.ai_validation_results, 
                                     st.session_state.ai_validation_output,
                                     st.session_state.ai_validation_filename,
                                     "AI Validation")
    
    with general_tab:
        st.markdown("### 📊 General Excel Comparison")
        st.markdown("**New Feature:** Compare any two Excel files with intelligent matching")
        
        # File upload section for general comparison
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown('<div class="upload-section">', unsafe_allow_html=True)
            st.markdown("### 📄 Source Excel File")
            st.markdown("Upload the first Excel file for comparison")
            
            source_excel = st.file_uploader(
                "Choose source Excel file",
                type=['xlsx', 'xls'],
                key="source_excel",
                help="Upload the first Excel file you want to compare"
            )
            
            if source_excel:
                st.markdown('<div class="upload-success">', unsafe_allow_html=True)
                st.markdown(f"### ✅ **{source_excel.name}**")
                st.markdown("📄 *Source file ready*")
                st.markdown('</div>', unsafe_allow_html=True)
                
                # Sheet selection for source
                if source_excel:
                    try:
                        import pandas as pd
                        xl_file = pd.ExcelFile(source_excel)
                        source_sheet = st.selectbox(
                            "Select source sheet:",
                            xl_file.sheet_names,
                            key="source_sheet"
                        )
                    except Exception as e:
                        st.error(f"Error reading file: {str(e)}")
            
            st.markdown('</div>', unsafe_allow_html=True)
        
        with col2:
            st.markdown('<div class="upload-section">', unsafe_allow_html=True)
            st.markdown("### 📄 Target Excel File")
            st.markdown("Upload the second Excel file for comparison")
            
            target_excel = st.file_uploader(
                "Choose target Excel file",
                type=['xlsx', 'xls'],
                key="target_excel",
                help="Upload the second Excel file you want to compare against"
            )
            
            if target_excel:
                st.markdown('<div class="upload-success">', unsafe_allow_html=True)
                st.markdown(f"### ✅ **{target_excel.name}**")
                st.markdown("📄 *Target file ready*")
                st.markdown('</div>', unsafe_allow_html=True)
                
                # Sheet selection for target
                if target_excel:
                    try:
                        import pandas as pd
                        xl_file = pd.ExcelFile(target_excel)
                        target_sheet = st.selectbox(
                            "Select target sheet:",
                            xl_file.sheet_names,
                            key="target_sheet"
                        )
                    except Exception as e:
                        st.error(f"Error reading file: {str(e)}")
            
            st.markdown('</div>', unsafe_allow_html=True)
        
        # General comparison configuration
        st.markdown("## ⚙️ General Comparison Settings")
        
        gen_col1, gen_col2, gen_col3, gen_col4 = st.columns(4)
        
        with gen_col1:
            gen_fuzzy_threshold = st.slider(
                "🎯 Matching Sensitivity",
                min_value=50,
                max_value=100,
                value=85,
                step=5,
                help="How closely data must match to be considered the same",
                key="gen_fuzzy_threshold"
            )
        
        with gen_col2:
            gen_max_results = st.number_input(
                "📊 Max Results",
                min_value=10,
                max_value=1000,
                value=200,
                step=10,
                key="gen_max_results"
            )
        
        with gen_col3:
            comparison_mode = st.selectbox(
                "🔍 Comparison Mode",
                ["Find Matches", "Find Differences", "Both"],
                help="What type of comparison to perform",
                key="comparison_mode"
            )
        
        with gen_col4:
            case_sensitive = st.checkbox(
                "🔤 Case Sensitive",
                value=False,
                help="Whether to consider letter case in comparisons",
                key="case_sensitive"
            )
        
        # Column mapping section
        if source_excel and target_excel:
            st.markdown("### 🔗 Column Mapping")
            st.markdown("Map columns between the two Excel files for comparison:")
            
            try:
                # Read sample data to show available columns
                source_df = pd.read_excel(source_excel, sheet_name=source_sheet if 'source_sheet' in locals() else 0, nrows=5)
                target_df = pd.read_excel(target_excel, sheet_name=target_sheet if 'target_sheet' in locals() else 0, nrows=5)
                
                col_map_col1, col_map_col2 = st.columns(2)
                
                with col_map_col1:
                    st.markdown("**Source Columns:**")
                    source_columns = source_df.columns.tolist()
                    for col in source_columns[:10]:  # Show first 10 columns
                        st.text(f"• {col}")
                    if len(source_columns) > 10:
                        st.text(f"... and {len(source_columns) - 10} more")
                
                with col_map_col2:
                    st.markdown("**Target Columns:**")
                    target_columns = target_df.columns.tolist()
                    for col in target_columns[:10]:  # Show first 10 columns
                        st.text(f"• {col}")
                    if len(target_columns) > 10:
                        st.text(f"... and {len(target_columns) - 10} more")
                
                # Key column selection
                key_col1, key_col2 = st.columns(2)
                with key_col1:
                    source_key_column = st.selectbox(
                        "Source key column:",
                        source_columns,
                        help="Main column to use for matching records",
                        key="source_key_column"
                    )
                
                with key_col2:
                    target_key_column = st.selectbox(
                        "Target key column:",
                        target_columns,
                        help="Main column to use for matching records",
                        key="target_key_column"
                    )
                
            except Exception as e:
                st.error(f"Error reading Excel files: {str(e)}")
        
        # Processing section for general comparison
        if source_excel and target_excel:
            if st.button("🚀 Start Excel Comparison", type="primary", key="start_general_comparison"):
                try:
                    with st.spinner("📊 Comparing Excel files..."):
                        results = process_general_comparison(
                            source_excel, target_excel,
                            source_sheet if 'source_sheet' in locals() else 0,
                            target_sheet if 'target_sheet' in locals() else 0,
                            source_key_column if 'source_key_column' in locals() else None,
                            target_key_column if 'target_key_column' in locals() else None,
                            gen_fuzzy_threshold, gen_max_results,
                            comparison_mode, case_sensitive
                        )
                    
                    if results:
                        st.session_state.general_comparison_results = results
                        st.success("✅ Excel comparison completed!")
                        st.rerun()
                    else:
                        st.error("❌ Comparison failed. Please check your files and try again.")
                        
                except Exception as e:
                    st.error(f"❌ Error during comparison: {str(e)}")
        else:
            st.info("📤 Upload both Excel files to begin comparison")
        
        # Display general comparison results
        if hasattr(st.session_state, 'general_comparison_results') and st.session_state.general_comparison_results:
            display_general_comparison_results(st.session_state.general_comparison_results)

def show_all_in_one():
    """All-in-One automated processor page - Complete workflow automation"""
    # Home button at the top
    col1, col2, col3 = st.columns([1, 8, 1])
    with col1:
        if st.button("🏠 Home", key="home_all_in_one"):
            st.session_state.current_page = "Dashboard"
            st.rerun()
    
    st.markdown('<h1 class="main-header"><span class="emoji-icon">⚡</span>All-in-One Student Opt-In Data Processor</h1>', unsafe_allow_html=True)
    
    # Workflow overview
    st.markdown("## 🔄 Automated Workflow Pipeline")
    
    # Dynamic workflow steps based on configuration
    if st.session_state.get('all_generate_traversa', True):
        col1, col2, col3, col4, col5 = st.columns(5)
        with col1:
            st.markdown("### 1️⃣ Word → PDF")
            st.markdown("📄 Convert documents")
        with col2:
            st.markdown("### 2️⃣ AI Extraction")
            st.markdown("🤖 Extract data")
        with col3:
            st.markdown("### 3️⃣ Validation")
            st.markdown("🔍 Compare data")
        with col4:
            st.markdown("### 4️⃣ Results")
            st.markdown("📊 Generate reports")
        with col5:
            st.markdown("### 5️⃣ Traversa")
            st.markdown("🚌 Routing ready")
    else:
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.markdown("### 1️⃣ Word → PDF")
            st.markdown("📄 Convert documents")
        with col2:
            st.markdown("### 2️⃣ AI Extraction")
            st.markdown("🤖 Extract data")
        with col3:
            st.markdown("### 3️⃣ Validation")
            st.markdown("🔍 Compare data")
        with col4:
            st.markdown("### 4️⃣ Results")
            st.markdown("📊 Generate reports")
    
    st.markdown("---")
    
    # File upload section
    st.markdown("## 📁 Upload Files")
    
    upload_col1, upload_col2 = st.columns(2)
    
    with upload_col1:
        st.markdown('<div class="upload-section">', unsafe_allow_html=True)
        st.markdown("### 📄 Word Documents")
        st.markdown("Upload student forms to process")
        
        word_files = st.file_uploader(
            "Choose Word documents",
            type=['docx', 'DOCX'],
            accept_multiple_files=True,
            key="all_word_files",
            help="Upload .docx files containing student forms"
        )
        
        if word_files:
            st.markdown('<div class="upload-success">', unsafe_allow_html=True)
            st.markdown(f"### ✅ **{len(word_files)} files uploaded**")
            st.markdown("📄 *Ready for processing*")
            st.markdown('</div>', unsafe_allow_html=True)
            
            # Show file list
            with st.expander("📋 View uploaded files"):
                for i, file in enumerate(word_files, 1):
                    st.markdown(f"{i}. {file.name} ({file.size / 1024:.1f} KB)")
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    with upload_col2:
        st.markdown('<div class="upload-section">', unsafe_allow_html=True)
        st.markdown("### 📊 Comparison Data")
        st.markdown("Upload student database for validation")
        
        comparison_file = st.file_uploader(
            "Choose comparison Excel file",
            type=['xlsx', 'xls'],
            key="all_comparison_file",
            help="Upload current student database export for validation"
        )
        
        if comparison_file:
            st.markdown('<div class="upload-success">', unsafe_allow_html=True)
            st.markdown(f"### ✅ **{comparison_file.name}**")
            st.markdown("📊 *Ready for validation*")
            st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Configuration section
    st.markdown("## ⚙️ Processing Configuration")
    
    config_col1, config_col2, config_col3, config_col4 = st.columns(4)
    
    with config_col1:
        st.markdown("#### 🤖 AI Model Settings")
        
        # Get available models
        available_models = {
            "Auto-detect (Recommended)": "auto",
            "ROCORI Transportation Forms": "rocorioptin", 
            "General Document": "prebuilt-document",
            "Daycare Forms": "daycareoptin2"
        }
        
        # Check for custom models
        if 'custom_models' in st.session_state:
            for name, model_id in st.session_state.custom_models.items():
                available_models[f"Custom: {name}"] = model_id
        
        ai_model = st.selectbox(
            "AI Model:",
            list(available_models.keys()),
            help="Select AI model for data extraction",
            key="all_ai_model"
        )
        
        ai_confidence = st.slider(
            "Confidence Threshold:",
            0.5, 1.0, 0.7, 0.05,
            help="Minimum confidence for AI extraction",
            key="all_ai_confidence"
        )
    
    with config_col2:
        st.markdown("#### 🔍 Validation Settings")
        
        matching_sensitivity = st.slider(
            "Matching Sensitivity:",
            50, 100, 80, 5,
            help="How closely data must match",
            key="all_matching_sensitivity"
        )
        
        max_results = st.number_input(
            "Max Results:",
            10, 1000, 100, 10,
            help="Maximum validation results",
            key="all_max_results"
        )
    
    with config_col3:
        st.markdown("#### 📄 Processing Options")
        
        processing_mode = st.selectbox(
            "Processing Mode:",
            ["Standard", "High Quality", "Fast"],
            help="Choose processing speed vs quality",
            key="all_processing_mode"
        )
        
        extract_options = st.multiselect(
            "Extract:",
            ["Tables", "Form Fields", "Full Text"],
            default=["Tables", "Form Fields"],
            help="What to extract from PDFs",
            key="all_extract_options"
        )
    
    with config_col4:
        st.markdown("#### 📊 Output Settings")
        
        include_intermediates = st.checkbox(
            "Save Intermediate Files",
            value=True,
            help="Save PDF and Excel files from each step",
            key="all_include_intermediates"
        )
        
        generate_traversa = st.checkbox(
            "Generate Traversa-Ready File",
            value=True,
            help="Create file ready for Traversa routing software upload",
            key="all_generate_traversa"
        )
        
        error_handling = st.selectbox(
            "Error Handling:",
            ["Continue Processing", "Stop on Error", "Skip Errors"],
            help="How to handle processing errors",
            key="all_error_handling"
        )
    
    # Processing section
    st.markdown("## 🚀 Start Automated Processing")
    
    if word_files and comparison_file:
        # Estimation
        total_files = len(word_files)
        estimated_time = total_files * 15  # Rough estimate: 15 seconds per file for full workflow
        
        st.info(f"⏱️ Estimated processing time: ~{estimated_time} seconds for {total_files} files")
        
        # Show workflow preview
        with st.expander("🔍 Workflow Preview"):
            traversa_step = f"""
            
            5. **Traversa Preparation** (Ready for routing software)
               - Remove unmatched students 
               - Update matched students with AI data
               - Highlight changes with color coding
               - Generate Traversa-ready Excel file""" if st.session_state.get('all_generate_traversa', True) else ""
            
            st.markdown(f"""
            **Processing Pipeline for {total_files} files:**
            
            1. **Word to PDF Conversion** ({total_files} files)
               - Convert each Word document to PDF format
               - Validate PDF generation
            
            2. **AI Data Extraction** ({total_files} PDFs)
               - Use {ai_model} model
               - Extract: {', '.join(extract_options)}
               - Confidence threshold: {ai_confidence:.0%}
            
            3. **Data Validation** (AI data vs {comparison_file.name})
               - Matching sensitivity: {matching_sensitivity}%
               - Maximum results: {max_results}
               - Error handling: {error_handling}
            
            4. **Results Generation**
               - Comprehensive validation report
               - Individual file processing logs
               - {"Intermediate files included" if include_intermediates else "Final results only"}{traversa_step}
            """)
        
        if st.button("🚀 Start Complete Processing", type="primary", key="start_all_processing"):
            st.session_state.all_in_one_processing = True
            st.rerun()
    else:
        missing_files = []
        if not word_files:
            missing_files.append("Word documents")
        if not comparison_file:
            missing_files.append("Comparison Excel file")
        
        st.warning(f"📤 Please upload: {', '.join(missing_files)}")
    
    # Processing execution
    if st.session_state.get('all_in_one_processing', False):
        execute_automated_workflow(
            word_files, comparison_file, 
            available_models[ai_model], ai_confidence,
            matching_sensitivity, max_results,
            processing_mode, extract_options,
            include_intermediates, error_handling,
            st.session_state.get('all_generate_traversa', True)
        )


def show_traversa_preparation():
    """Traversa data preparation page - Format data specifically for Traversa routing software"""
    # Home button at the top
    col1, col2, col3 = st.columns([1, 8, 1])
    with col1:
        if st.button("🏠 Home", key="home_traversa_prep"):
            st.session_state.current_page = "Dashboard"
            st.rerun()
    
    # Header below the home button for better alignment
    st.markdown('<h1 class="main-header"><span class="emoji-icon">🚌</span>Traversa Data Preparation</h1>', unsafe_allow_html=True)
    
    st.markdown('<p class="subtitle">Professional transportation data formatting for Traversa routing software • Voigts Excellence in Action</p>', unsafe_allow_html=True)
    
    if not TRAVERSA_PROCESSOR_AVAILABLE:
        st.error("❌ Traversa data processor is not available. Please check the installation.")
        return
    
    st.markdown("""
    ### 🚌 About Traversa Preparation
    
    This tool processes your student data comparison results specifically for Traversa routing software:
    
    **Key Features:**
    - 📋 **Maintains Template Format** - Keeps your original Traversa import template structure
    - ❌ **Removes Unmatched Students** - Only keeps students that were successfully matched
    - 🔄 **Updates Information** - Replaces old data with AI-extracted information where matches exist
    - 🎨 **Highlights Changes** - Shows exactly what information was updated from the original
    - ✅ **Traversa Ready** - Output file is ready for direct upload to Traversa routing software
    
    **Process Flow:**
    1. Upload your AI-extracted student data (Excel file)
    2. Upload your Traversa template/existing data (Excel file)
    3. Configure field mappings (auto-detected or manual)
    4. Download the Traversa-ready file with highlighted changes
    """)
    
    st.markdown("---")
    
    # File upload sections
    st.markdown("## 📂 Upload Your Files")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown('<div class="upload-section">', unsafe_allow_html=True)
        st.markdown("### 🤖 AI Extracted Data")
        st.markdown("Upload the Excel file containing AI-extracted student data")
        
        ai_file = st.file_uploader(
            "Choose AI extracted data file",
            type=['xlsx', 'xls'],
            key="traversa_ai_file",
            help="Excel file with student data extracted by AI from PDF forms"
        )
        
        if ai_file:
            st.success(f"✅ AI file loaded: {ai_file.name}")
            
            # Show preview of AI data
            try:
                df_preview = pd.read_excel(ai_file, nrows=3)
                st.markdown("**Preview:**")
                st.dataframe(df_preview, use_container_width=True)
                st.info(f"📊 Contains {len(pd.read_excel(ai_file))} students with {len(df_preview.columns)} fields")
            except Exception as e:
                st.warning(f"⚠️ Could not preview file: {e}")
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="upload-section">', unsafe_allow_html=True)
        st.markdown("### 📋 Traversa Template")
        st.markdown("Upload your Traversa import template or existing student data")
        
        traversa_file = st.file_uploader(
            "Choose Traversa template file",
            type=['xlsx', 'xls'],
            key="traversa_template_file",
            help="Excel file in Traversa import format with existing student data"
        )
        
        if traversa_file:
            st.success(f"✅ Traversa template loaded: {traversa_file.name}")
            
            # Show preview of Traversa template
            try:
                df_preview = pd.read_excel(traversa_file, nrows=3)
                st.markdown("**Preview:**")
                st.dataframe(df_preview, use_container_width=True)
                st.info(f"📊 Contains {len(pd.read_excel(traversa_file))} students with {len(df_preview.columns)} fields")
            except Exception as e:
                st.warning(f"⚠️ Could not preview file: {e}")
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Processing configuration
    if ai_file and traversa_file:
        st.markdown("---")
        st.markdown("## ⚙️ Processing Configuration")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("### 🔧 Matching Settings")
            
            fuzzy_threshold = st.slider(
                "Name Matching Sensitivity",
                min_value=60,
                max_value=100,
                value=80,
                help="Higher values require more exact name matches (recommended: 80)"
            )
            
            auto_map_fields = st.checkbox(
                "Auto-map similar field names",
                value=True,
                help="Automatically detect and map similar column names between files"
            )
        
        with col2:
            st.markdown("### 📋 Field Mapping")
            
            if not auto_map_fields:
                st.info("💡 Manual field mapping - specify how AI fields map to Traversa columns")
                
                # Get column names for manual mapping
                try:
                    ai_columns = list(pd.read_excel(ai_file, nrows=0).columns)
                    traversa_columns = list(pd.read_excel(traversa_file, nrows=0).columns)
                    
                    st.markdown("**Available AI Fields:** " + ", ".join(ai_columns[:5]) + ("..." if len(ai_columns) > 5 else ""))
                    st.markdown("**Available Traversa Fields:** " + ", ".join(traversa_columns[:5]) + ("..." if len(traversa_columns) > 5 else ""))
                    
                    mapping_text = st.text_area(
                        "Field Mappings (AI_field:Traversa_field, one per line)",
                        placeholder="Student Name:Student_Name\nGrade:Grade_Level\nAddress:Home_Address",
                        help="Map AI extracted fields to Traversa template columns"
                    )
                except Exception as e:
                    st.warning(f"⚠️ Could not load column information: {e}")
                    mapping_text = ""
            else:
                st.info("🔍 Field mapping will be automatically detected based on column name similarity")
                mapping_text = ""
        
        # Process button
        st.markdown("---")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("🚌 Prepare Data for Traversa", use_container_width=True, type="primary"):
                with st.spinner("🔄 Processing data for Traversa..."):
                    try:
                        # Create temporary files
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as ai_temp:
                            ai_temp.write(ai_file.getvalue())
                            ai_temp_path = ai_temp.name
                        
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as traversa_temp:
                            traversa_temp.write(traversa_file.getvalue())
                            traversa_temp_path = traversa_temp.name
                        if not TRAVERSA_PROCESSOR_AVAILABLE:
                            st.error("❌ Traversa processor not available in this deployment")
                            st.info("This feature requires additional modules not included in the cloud deployment")
                            return
                        
                        # Initialize processor
                        processor = TraversaDataProcessor()
                        
                        # Set manual field mappings if provided
                        if not auto_map_fields and mapping_text:
                            mappings = {}
                            for line in mapping_text.strip().split('\n'):
                                if ':' in line:
                                    ai_field, traversa_field = line.split(':', 1)
                                    mappings[ai_field.strip()] = traversa_field.strip()
                            processor.set_field_mappings(mappings)
                        
                        # Create output file path
                        output_filename = f"traversa_ready_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                        
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as output_temp:
                            output_temp_path = output_temp.name
                        
                        # Process data
                        results = processor.process_for_traversa(
                            ai_extractor_file=ai_temp_path,
                            traversa_template_file=traversa_temp_path,
                            output_file=output_temp_path,
                            fuzzy_threshold=fuzzy_threshold,
                            auto_map_fields=auto_map_fields
                        )
                        
                        # Display results
                        st.success("✅ Traversa data preparation completed!")
                        
                        # Results summary
                        st.markdown("### 📊 Processing Results")
                        
                        col1, col2, col3, col4 = st.columns(4)
                        
                        with col1:
                            st.markdown('<div class="metric-card">', unsafe_allow_html=True)
                            st.metric("Original Students", results['total_original_students'])
                            st.markdown('</div>', unsafe_allow_html=True)
                        
                        with col2:
                            st.markdown('<div class="metric-card">', unsafe_allow_html=True)
                            st.metric("Matched & Kept", results['matched_students'])
                            st.markdown('</div>', unsafe_allow_html=True)
                        
                        with col3:
                            st.markdown('<div class="metric-card">', unsafe_allow_html=True)
                            st.metric("Removed", results['removed_students'])
                            st.markdown('</div>', unsafe_allow_html=True)
                        
                        with col4:
                            st.markdown('<div class="metric-card">', unsafe_allow_html=True)
                            st.metric("Fields Updated", results['updated_fields'])
                            st.markdown('</div>', unsafe_allow_html=True)
                        
                        # Field mappings used
                        if results['field_mappings_used']:
                            st.markdown("### 🔗 Field Mappings Applied")
                            mapping_df = pd.DataFrame([
                                {"AI Field": k, "Traversa Field": v} 
                                for k, v in results['field_mappings_used'].items()
                            ])
                            st.dataframe(mapping_df, use_container_width=True)
                        
                        # Changes summary
                        if results['changes_summary'] and results['changes_summary'].get('by_category'):
                            st.markdown("### 🎨 Changes by Category")
                            
                            # Category summary with color indicators
                            category_colors = {
                                'address': '🔵 Address (Blue)',
                                'student_name': '🟣 Student Name (Purple)', 
                                'daycare': '🟢 Daycare (Green)',
                                'general': '� General (Orange)'
                            }
                            
                            category_df = pd.DataFrame([
                                {"Category": category_colors.get(k, k.title()), "Updates": v} 
                                for k, v in results['changes_summary']['by_category'].items()
                            ])
                            st.dataframe(category_df, use_container_width=True)
                            
                            # Field-level changes
                            if results['changes_summary'].get('by_field'):
                                st.markdown("### �🔄 Changes by Field")
                                changes_df = pd.DataFrame([
                                    {"Field": k, "Updates": v} 
                                    for k, v in results['changes_summary']['by_field'].items()
                                ])
                                st.dataframe(changes_df, use_container_width=True)
                        
                        # Download button
                        st.markdown("### 📥 Download Traversa-Ready File")
                        
                        with open(output_temp_path, 'rb') as f:
                            file_data = f.read()
                        
                        st.download_button(
                            label="📥 Download Traversa-Ready Excel File",
                            data=file_data,
                            file_name=output_filename,
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                        
                        st.markdown("""
                        ### ✅ Next Steps
                        1. **Download** the Traversa-ready file above
                        2. **Review** the highlighted changes in the Excel file
                        3. **Upload** directly to Traversa routing software
                        4. The file maintains your original template format for compatibility
                        """)
                        
                        # Cleanup temporary files
                        try:
                            os.unlink(ai_temp_path)
                            os.unlink(traversa_temp_path)
                            os.unlink(output_temp_path)
                        except:
                            pass
                            
                    except Exception as e:
                        st.error(f"❌ Error processing data: {e}")
                        st.error("Please check your files and try again.")


def main():
    """Main application function with navigation"""
    
    initialize_session_state()
    
    # Create navigation menu and get selected page
    current_page = create_navigation_menu()
    
    # Custom Logo Section (shown on all pages) - Perfect centering
    logo_path = "Voigts Bus Service Logo.png"
    
    if os.path.exists(logo_path):
        # Read image and encode as base64 for inline display
        import base64
        with open(logo_path, "rb") as img_file:
            img_base64 = base64.b64encode(img_file.read()).decode()
        
        # Use Streamlit columns for perfect centering
        col1, col2, col3 = st.columns([1, 1, 1])
        with col2:
            st.markdown(f"""
            <div style="
                text-align: center !important;
                display: flex !important;
                justify-content: center !important;
                align-items: center !important;
                width: 100% !important;
            ">
                <img src="data:image/png;base64,{img_base64}" 
                     style="
                         width: 150px !important;
                         height: auto !important;
                         display: block !important;
                         margin: 0 auto !important;
                         object-fit: contain !important;
                     ">
            </div>
            """, unsafe_allow_html=True)
    else:
        # Use Streamlit columns for emoji fallback too
        col1, col2, col3 = st.columns([1, 1, 1])
        with col2:
            st.markdown("""
            <div style="
                text-align: center !important;
                font-size: 3rem;
                display: flex !important;
                justify-content: center !important;
                align-items: center !important;
                width: 100% !important;
            ">🚌</div>
            """, unsafe_allow_html=True)
    
    st.markdown("""
    <div class="logo-section" style="text-align: center; margin-bottom: 20px; margin-top: 5px;">
        <h2>Voigt's Bus Companies <span style="color: #1e3c72;">✕</span> Chayton Creations Co.</h2>
        <p style="color: #1e3c72; margin: 0; font-weight: 500;">Student Opt-In Data Management Solutions</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Display selected page
    if current_page == "Dashboard":
        show_dashboard()
    elif current_page == "Word to PDF":
        show_word_to_pdf()
    elif current_page == "PDF Extraction":
        show_pdf_extraction()
    elif current_page == "Data Validation":
        show_data_validation()
    elif current_page == "Traversa Preparation":
        show_traversa_preparation()
    elif current_page == "All-in-One":
        show_all_in_one()
    
    # Footer (shown on all pages)
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: #333; padding: 20px;'>
        <p><strong>Voigt's Bus Companies X Chayton Creations Co.</strong> | Student Opt-In Data Management</p>
        <p>Made with 💙 for | <strong>Sauk Rapids-Rice</strong> | <strong>ROCORI</strong> | <strong>Stride Academy</strong> | <strong>Math & Science Academy</strong></p>
        <p>🤖 Powered by advanced AI technology for seamless transportation data processing</p>
    </div>
    """, unsafe_allow_html=True)

# Helper functions for data validation (simplified versions)
def process_comparison(ai_file, comparison_file, fuzzy_threshold, max_results):
    """Process comparison between AI extracted data and comparison file"""
    try:
        # Import the comparison functionality
        from student_data_comparator import StudentDataComparator
        
        # Create comparator instance
        comparator = StudentDataComparator(
            ai_file=ai_file,
            comparison_file=comparison_file,
            fuzzy_threshold=fuzzy_threshold,
            max_results=max_results
        )
        
        # Run comparison
        results = comparator.run_comparison()
        
        # Generate output file
        output_data, output_filename = comparator.generate_output()
        
        return results, output_data, output_filename
        
    except Exception as e:
        st.error(f"Error during comparison: {str(e)}")
        return None, None, None

def display_validation_results(results, output_data, output_filename, title="Validation Results"):
    """Display validation results from AI data comparison"""
    st.markdown(f"## 📊 {title}")
    
    # Summary metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total Records", len(results))
    with col2:
        matched = len([r for r in results if r.get('match_found', False)])
        st.metric("Matches Found", matched)
    with col3:
        unmatched = len(results) - matched
        st.metric("No Matches", unmatched)
    with col4:
        match_rate = (matched / len(results)) * 100 if results else 0
        st.metric("Match Rate", f"{match_rate:.1f}%")
    
    # Download section
    st.markdown("### 💾 Download Results")
    download_link = create_download_link(output_data, output_filename, "📥 Download Validation Results")
    st.markdown(download_link, unsafe_allow_html=True)
    
    # Detailed results preview
    with st.expander("📋 View Detailed Results"):
        for i, result in enumerate(results[:10]):  # Show first 10 results
            st.markdown(f"**Record {i+1}:** {result}")

def display_general_comparison_results(results):
    """Display results from general Excel comparison"""
    st.markdown("## 📊 Excel Comparison Results")
    
    # Results summary
    total_matches = results.get('matches', [])
    total_differences = results.get('differences', [])
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Matches Found", len(total_matches))
    with col2:
        st.metric("Differences Found", len(total_differences))
    with col3:
        total_compared = len(total_matches) + len(total_differences)
        st.metric("Total Compared", total_compared)
    
    # Display matches
    if total_matches:
        with st.expander(f"✅ Matches ({len(total_matches)})"):
            for i, match in enumerate(total_matches[:20]):  # Show first 20
                st.markdown(f"**{i+1}.** {match}")
    
    # Display differences
    if total_differences:
        with st.expander(f"❌ Differences ({len(total_differences)})"):
            for i, diff in enumerate(total_differences[:20]):  # Show first 20
                st.markdown(f"**{i+1}.** {diff}")

def process_general_comparison(source_excel, target_excel, source_sheet, target_sheet, 
                             source_key_col, target_key_col, fuzzy_threshold, max_results,
                             comparison_mode, case_sensitive):
    """Process general Excel file comparison"""
    try:
        import pandas as pd
        
        # Read the Excel files
        source_df = pd.read_excel(source_excel, sheet_name=source_sheet)
        target_df = pd.read_excel(target_excel, sheet_name=target_sheet)
        
        matches = []
        differences = []
        
        # Simple comparison logic (can be enhanced)
        if source_key_col and target_key_col:
            # Key-based comparison
            for idx, source_row in source_df.iterrows():
                source_key = source_row[source_key_col]
                if not case_sensitive:
                    source_key = str(source_key).lower()
                
                # Look for matches in target
                target_matches = target_df[target_df[target_key_col].astype(str).str.lower() == str(source_key).lower()]
                
                if not target_matches.empty:
                    matches.append(f"{source_key} found in both files")
                else:
                    differences.append(f"{source_key} only in source file")
        else:
            # Simple row-by-row comparison
            min_rows = min(len(source_df), len(target_df))
            for i in range(min_rows):
                source_row = source_df.iloc[i].to_string()
                target_row = target_df.iloc[i].to_string()
                
                if source_row == target_row:
                    matches.append(f"Row {i+1} matches")
                else:
                    differences.append(f"Row {i+1} differs")
        
        return {
            'matches': matches[:max_results],
            'differences': differences[:max_results],
            'source_rows': len(source_df),
            'target_rows': len(target_df)
        }
        
    except Exception as e:
        st.error(f"Error during comparison: {str(e)}")
        return None

def execute_automated_workflow(word_files, comparison_file, ai_model_id, ai_confidence,
                             matching_sensitivity, max_results, processing_mode, 
                             extract_options, include_intermediates, error_handling, generate_traversa=True):
    """Execute the complete automated workflow"""
    
    st.markdown("## 🚀 Processing in Progress")
    
    # Initialize workflow state
    workflow_state = {
        'total_files': len(word_files),
        'current_step': 'Starting',
        'processed_files': 0,
        'errors': [],
        'results': {}
    }
    
    # Create containers for dynamic updates
    main_progress = st.progress(0)
    status_container = st.empty()
    step_container = st.container()
    
    def update_status(message, progress, step="Processing"):
        workflow_state['current_step'] = step
        main_progress.progress(progress)
        status_container.text(f"🔄 {step}: {message}")
    
    try:
        # STEP 1: Word to PDF Conversion
        update_status("Converting Word documents to PDF...", 0.1, "Step 1/4")
        
        with step_container:
            st.markdown("### 📄 Step 1: Word to PDF Conversion")
            pdf_progress = st.progress(0)
            pdf_status = st.empty()
        
        pdf_files = []
        pdf_data = {}
        
        with tempfile.TemporaryDirectory() as temp_dir:
            for i, word_file in enumerate(word_files):
                try:
                    pdf_status.text(f"Converting {word_file.name}...")
                    
                    # Convert Word to PDF
                    pdf_path, pdf_filename = convert_docx_to_pdf(word_file, temp_dir)
                    
                    # Read PDF data
                    with open(pdf_path, 'rb') as f:
                        pdf_data[pdf_filename] = f.read()
                    
                    pdf_files.append({
                        'name': pdf_filename,
                        'data': pdf_data[pdf_filename],
                        'source': word_file.name
                    })
                    
                    workflow_state['processed_files'] += 1
                    pdf_progress.progress((i + 1) / len(word_files))
                    
                except Exception as e:
                    error_msg = f"Failed to convert {word_file.name}: {str(e)}"
                    workflow_state['errors'].append(error_msg)
                    pdf_status.error(f"❌ {error_msg}")
                    
                    if error_handling == "Stop on Error":
                        raise Exception(error_msg)
                    elif error_handling == "Continue Processing":
                        continue
            
            pdf_status.success(f"✅ Converted {len(pdf_files)} files to PDF")
            
            # STEP 2: AI Data Extraction
            update_status("Extracting data using AI...", 0.4, "Step 2/4")
            
            with step_container:
                st.markdown("### 🤖 Step 2: AI Data Extraction")
                ai_progress = st.progress(0)
                ai_status = st.empty()
            
            # Prepare PDF files for AI extraction
            ai_files = []
            for pdf_file in pdf_files:
                # Create file-like object from bytes
                from io import BytesIO
                pdf_bytes = BytesIO(pdf_file['data'])
                pdf_bytes.name = pdf_file['name']
                ai_files.append(pdf_bytes)
            
            def ai_progress_callback(message, progress):
                ai_status.text(f"🤖 {message}")
                ai_progress.progress(progress)
            
            # Extract data using AI
            extract_options_dict = {
                'tables': 'Tables' in extract_options,
                'forms': 'Form Fields' in extract_options,
                'text': 'Full Text' in extract_options,
                'confidence_threshold': ai_confidence
            }
            
            excel_data, ai_filename, extracted_data = extract_data_from_pdfs(
                ai_files,
                progress_callback=ai_progress_callback,
                model_id=ai_model_id,
                extract_options=extract_options_dict
            )
            
            ai_status.success(f"✅ Extracted data from {len(pdf_files)} files")
            
            # Store AI extraction results
            workflow_state['results']['ai_excel_data'] = excel_data
            workflow_state['results']['ai_filename'] = ai_filename
            workflow_state['results']['extracted_data'] = extracted_data
            
            # STEP 3: Data Validation
            update_status("Validating extracted data...", 0.7, "Step 3/4")
            
            with step_container:
                st.markdown("### 🔍 Step 3: Data Validation")
                val_progress = st.progress(0)
                val_status = st.empty()
            
            # Create AI file object for validation
            ai_file_obj = BytesIO(excel_data)
            ai_file_obj.name = ai_filename
            
            val_status.text("🔍 Comparing AI data with student database...")
            val_progress.progress(0.5)
            
            # Run validation
            validation_results, validation_output, validation_filename = process_comparison(
                ai_file_obj, comparison_file, matching_sensitivity, max_results
            )
            
            val_progress.progress(1.0)
            val_status.success("✅ Data validation completed")
            
            # Store validation results
            workflow_state['results']['validation_results'] = validation_results
            workflow_state['results']['validation_output'] = validation_output
            workflow_state['results']['validation_filename'] = validation_filename
            
            # STEP 5: Traversa Processing (Optional)
            traversa_step_num = "5/5" if generate_traversa else "4/4"
            
            if generate_traversa and TRAVERSA_PROCESSOR_AVAILABLE:
                update_status("Preparing Traversa-ready file...", 0.8, f"Step 5/{traversa_step_num}")
                
                with step_container:
                    st.markdown("### 🚌 Step 5: Traversa Data Preparation")
                    traversa_progress = st.progress(0)
                    traversa_status = st.empty()
                
                try:
                    traversa_status.text("🚌 Processing data for Traversa routing software...")
                    traversa_progress.progress(0.3)
                    
                    # Create temporary files for Traversa processing
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as ai_temp:
                        ai_temp.write(workflow_state['results']['ai_excel_data'])
                        ai_temp_path = ai_temp.name
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as comparison_temp:
                        comparison_temp.write(comparison_file.getvalue())
                        comparison_temp_path = comparison_temp.name
                    
                    if not TRAVERSA_PROCESSOR_AVAILABLE:
                        st.error("❌ Traversa processor not available in this deployment")
                        st.info("This feature requires additional modules not included in the cloud deployment")
                        return
                    
                    # Initialize Traversa processor
                    processor = TraversaDataProcessor()
                    
                    traversa_progress.progress(0.5)
                    traversa_status.text("🔄 Matching students and updating data...")
                    
                    # Create output file path
                    traversa_filename = f"traversa_ready_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
                    
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as traversa_output_temp:
                        traversa_output_path = traversa_output_temp.name
                    
                    # Process data for Traversa
                    traversa_results = processor.process_for_traversa(
                        ai_extractor_file=ai_temp_path,
                        traversa_template_file=comparison_temp_path,
                        output_file=traversa_output_path,
                        fuzzy_threshold=matching_sensitivity,
                        auto_map_fields=True
                    )
                    
                    traversa_progress.progress(0.8)
                    traversa_status.text("📊 Generating Traversa report...")
                    
                    # Read the Traversa output
                    with open(traversa_output_path, 'rb') as f:
                        traversa_data = f.read()
                    
                    # Store Traversa results
                    workflow_state['results']['traversa_data'] = traversa_data
                    workflow_state['results']['traversa_filename'] = traversa_filename
                    workflow_state['results']['traversa_results'] = traversa_results
                    
                    traversa_progress.progress(1.0)
                    traversa_status.success("✅ Traversa-ready file generated successfully!")
                    
                    # Cleanup temporary files
                    try:
                        os.unlink(ai_temp_path)
                        os.unlink(comparison_temp_path)
                        os.unlink(traversa_output_path)
                    except:
                        pass
                        
                except Exception as e:
                    traversa_status.error(f"❌ Traversa processing failed: {str(e)}")
                    workflow_state['errors'].append(f"Traversa processing failed: {str(e)}")
                    
                    if error_handling == "Stop on Error":
                        raise Exception(f"Traversa processing failed: {str(e)}")
            
            # STEP 4/5: Results Generation
            final_step_num = "5/5" if generate_traversa else "4/4"
            update_status("Generating comprehensive results...", 0.9, f"Step {final_step_num}")
            
            with step_container:
                step_title = "### 📊 Step 6: Results Generation" if generate_traversa else "### 📊 Step 4: Results Generation"
                st.markdown(step_title)
                results_progress = st.progress(0)
                results_status = st.empty()
            
            results_progress.progress(0.5)
            results_status.text("📊 Compiling workflow results...")
            
            # Generate comprehensive workflow report
            workflow_report = generate_workflow_report(
                workflow_state, pdf_files, extracted_data, 
                validation_results, include_intermediates
            )
            
            results_progress.progress(1.0)
            results_status.success("✅ Workflow completed successfully!")
            
            # Final status
            update_status("Workflow completed successfully!", 1.0, "Complete")
            
            # Display results
            display_workflow_results(workflow_state, workflow_report, include_intermediates)
            
    except Exception as e:
        st.error(f"❌ Workflow failed: {str(e)}")
        
        # Show error details
        with st.expander("🔍 Error Details"):
            st.text(str(e))
            if workflow_state['errors']:
                st.markdown("**Individual Errors:**")
                for error in workflow_state['errors']:
                    st.text(f"• {error}")
    
    finally:
        # Reset processing state
        st.session_state.all_in_one_processing = False

def generate_workflow_report(workflow_state, pdf_files, extracted_data, validation_results, include_intermediates):
    """Generate comprehensive workflow report"""
    
    report = {
        'summary': {
            'total_files': workflow_state['total_files'],
            'processed_files': workflow_state['processed_files'],
            'errors': len(workflow_state['errors']),
            'success_rate': (workflow_state['processed_files'] / workflow_state['total_files']) * 100 if workflow_state['total_files'] > 0 else 0
        },
        'conversion_results': pdf_files,
        'extraction_results': extracted_data,
        'validation_results': validation_results,
        'errors': workflow_state['errors'],
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    
    return report

def display_workflow_results(workflow_state, workflow_report, include_intermediates):
    """Display comprehensive workflow results"""
    
    st.markdown("## 🎉 Workflow Results")
    
    # Summary metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total Files", workflow_state['total_files'])
    with col2:
        st.metric("Processed", workflow_state['processed_files'])
    with col3:
        st.metric("Errors", len(workflow_state['errors']))
    with col4:
        success_rate = workflow_report['summary']['success_rate']
        st.metric("Success Rate", f"{success_rate:.1f}%")
    
    # Download section
    st.markdown("### 💾 Download Results")
    
    # Determine number of columns based on available results
    has_traversa = 'traversa_data' in workflow_state['results']
    
    if has_traversa:
        download_col1, download_col2, download_col3, download_col4 = st.columns(4)
    else:
        download_col1, download_col2, download_col3 = st.columns(3)
    
    with download_col1:
        # Main validation results
        if 'validation_output' in workflow_state['results']:
            st.download_button(
                label="📊 Download Validation Results",
                data=workflow_state['results']['validation_output'],
                file_name=workflow_state['results']['validation_filename'],
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                help="Main validation results Excel file"
            )
    
    with download_col2:
        # AI extraction results
        if 'ai_excel_data' in workflow_state['results']:
            st.download_button(
                label="🤖 Download AI Extracted Data",
                data=workflow_state['results']['ai_excel_data'],
                file_name=workflow_state['results']['ai_filename'],
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                help="Raw AI extraction results"
            )
    
    with download_col3:
        # Comprehensive workflow report
        import json
        report_json = json.dumps(workflow_report, indent=2, default=str)
        st.download_button(
            label="📋 Download Workflow Report",
            data=report_json,
            file_name=f"workflow_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mime="application/json",
            help="Complete workflow log and results"
        )
    
    if has_traversa:
        with download_col4:
            # Traversa-ready file
            st.download_button(
                label="🚌 Download Traversa-Ready File",
                data=workflow_state['results']['traversa_data'],
                file_name=workflow_state['results']['traversa_filename'],
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                help="Ready for Traversa routing software upload",
                type="primary"
            )
    
    # Detailed results sections
    with st.expander("📄 PDF Conversion Results"):
        if workflow_report['conversion_results']:
            for pdf in workflow_report['conversion_results']:
                st.markdown(f"✅ **{pdf['source']}** → {pdf['name']}")
        else:
            st.info("No PDF conversions completed")
    
    with st.expander("🤖 AI Extraction Results"):
        if workflow_report['extraction_results']:
            for result in workflow_report['extraction_results']:
                status = "✅" if 'error' not in result else "❌"
                st.markdown(f"{status} **{result.get('source_file', 'Unknown')}**")
                if 'error' in result:
                    st.text(f"   Error: {result['error']}")
                else:
                    st.text(f"   Confidence: {result.get('confidence', 0):.1%}")
        else:
            st.info("No AI extraction results available")
    
    with st.expander("🔍 Validation Summary"):
        if workflow_report['validation_results']:
            matched = len([r for r in workflow_report['validation_results'] if r.get('match_found', False)])
            total = len(workflow_report['validation_results'])
            st.markdown(f"**Matches Found:** {matched}/{total} ({matched/total*100:.1f}%)")
        else:
            st.info("No validation results available")
    
    # Traversa processing results (if available)
    if 'traversa_results' in workflow_state['results']:
        with st.expander("🚌 Traversa Processing Results"):
            traversa_results = workflow_state['results']['traversa_results']
            
            # Summary statistics
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Students Matched", traversa_results.get('matched_students', 0))
            with col2:
                st.metric("Students Removed", traversa_results.get('removed_students', 0))
            with col3:
                st.metric("Fields Updated", traversa_results.get('updated_fields', 0))
            
            # Changes by category
            if traversa_results.get('changes_summary', {}).get('by_category'):
                st.markdown("**Changes by Category:**")
                category_colors = {
                    'address': '🔵 Address Changes',
                    'student_name': '🟣 Student Name Changes', 
                    'daycare': '🟢 Daycare Changes',
                    'general': '🟡 General Changes'
                }
                
                for category, count in traversa_results['changes_summary']['by_category'].items():
                    display_name = category_colors.get(category, category.title())
                    st.markdown(f"   {display_name}: {count}")
            
            st.success("✅ File is ready for direct upload to Traversa routing software!")
    
    if workflow_state['errors']:
        with st.expander("❌ Error Log"):
            for error in workflow_state['errors']:
                st.text(f"• {error}")
    
    # Reset button
    st.markdown("---")
    if st.button("🔄 Start New Workflow"):
        # Clear all workflow results
        for key in ['all_in_one_processing', 'validation_results', 'extraction_results']:
            if key in st.session_state:
                del st.session_state[key]
        st.rerun()

if __name__ == "__main__":
    main()
